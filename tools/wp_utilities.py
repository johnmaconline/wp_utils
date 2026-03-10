##########################################################################################
#
# Script name: wp_utilities.py
#
# Description: Downloads a specified number of blog posts as text documents from a given URL.
#              Each blog post is saved as a text file with the blog title as the filename
#              in the specified output directory.
#
# Author: [Your Name]
#
##########################################################################################

import base64
import requests
from bs4 import BeautifulSoup, NavigableString
import argparse
import logging
import sys
import os
import json
import mimetypes
from datetime import date
from typing import Any
import atexit
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from datetime import datetime, timedelta, time
from urllib.parse import urlparse, unquote
from zoneinfo import ZoneInfo
from markdown import markdown as render_markdown

# ****************************************************************************************
# Global data and configuration
# ****************************************************************************************

# Logging config
log = logging.getLogger(os.path.basename(sys.argv[0]))
log.setLevel(logging.DEBUG)

# File handler for logging
fh = logging.FileHandler('wp_utilities.log', mode='w')
fh.setLevel(logging.DEBUG)
formatter = logging.Formatter(
    '%(asctime)-15s [%(funcName)25s:%(lineno)-5s] %(levelname)-8s %(message)s')
fh.setFormatter(formatter)
log.addHandler(fh)

# Number of consecutive filename-based duplicates before stopping early
NUM_DUPLICATES = 5
# WordPress API max per_page is typically 100
WP_API_MAX_PER_PAGE = 100
# ASCII table rendering
TABLE_MAX_WIDTH = 120
TABLE_MIN_WIDTH = 4
TABLE_MIN_WIDTH_TITLE = 20
TABLE_MIN_WIDTH_DATE = 10
# Export defaults
EXPORT_DEFAULT_DIR = 'export'
EXPORT_RESOURCES = {
    'posts',
    'pages',
    'media',
    'comments',
    'users',
    'categories',
    'tags',
    'taxonomies',
    'types',
    'statuses',
    'settings',
    'menus',
    'plugins',
}
DEFAULT_TIMEZONE = 'America/New_York'
POST_NAVIGATION_BLOCKS = """<!-- wp:columns -->
<div class="wp-block-columns"><!-- wp:column {"width":"25%"} -->
<div class="wp-block-column" style="flex-basis:25%"><!-- wp:post-navigation-link {"type":"previous"} /--></div>
<!-- /wp:column -->

<!-- wp:column {"width":"50%"} -->
<div class="wp-block-column" style="flex-basis:50%"></div>
<!-- /wp:column -->

<!-- wp:column {"width":"25%"} -->
<div class="wp-block-column" style="flex-basis:25%"><!-- wp:post-navigation-link {"type":"next"} /--></div>
<!-- /wp:column --></div>
<!-- /wp:columns -->"""
POST_NAVIGATION_BLOCKS_PATTERN = re.compile(
    r'<!-- wp:columns -->.*?wp:post-navigation-link.*?<!-- /wp:columns -->',
    re.DOTALL,
)

# ****************************************************************************************
# Exceptions
# ****************************************************************************************

class Error(Exception):
    """Base class for exceptions in this module."""
    pass

class RequestError(Error):
    def __init__(self, url):
        self.message = f"Failed to fetch URL: {url}"
        super().__init__(self.message)

# ****************************************************************************************
# Functions
# ****************************************************************************************
def valid_filename(url):
    """
    Create a valid filename from a URL by keeping only alphanumeric characters and replacing others,
    while preserving the original file extension.
    """
    parsed_url = urlparse(url)
    filename, file_extension = os.path.splitext(os.path.basename(parsed_url.path))
    # Keep alphanumeric characters and replace others with underscore
    filename = re.sub(r'\W+', '_', unquote(filename))
    # Concatenate filename and file extension
    return filename + file_extension


def download_image(image_url, outdir):
    '''
    Download an image and save it to the specified subdirectory ('images') within outdir.
    Returns the local filename of the downloaded image.
    '''
    images_dir = os.path.join(outdir, 'images')
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)

    # Extract base URL (without query parameters) and create a valid filename
    base_url = image_url.split('?')[0]
    local_filename = os.path.join(images_dir, valid_filename(base_url))

    try:
        response = requests.get(image_url, stream=True)
        if response.status_code == 200:
            with open(local_filename, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            # Verify the downloaded file
            if os.path.getsize(local_filename) > 0:
                return local_filename
            else:
                log.error(f"Downloaded image is empty: {image_url}")
        else:
            log.error(f"Failed to download image: {image_url}, Status code: {response.status_code}")
    except Exception as e:
        log.error(f"Error downloading image {image_url}: {e}")
        
    log.debug(f"Image downloaded and saved to: {local_filename}")


    return None


def extract_text_from_html(html_content, outdir):
    '''
    Extract text from HTML content (WordPress API) including headings, lists, code blocks, and images.
    '''
    soup = BeautifulSoup(html_content or "", 'html.parser')
    container = soup.body if soup.body else soup
    content_text = ''

    def append_code_block(code_text):
        nonlocal content_text
        if not code_text:
            return
        code_text = code_text.strip('\n')
        if not code_text.strip():
            return
        content_text += f"```\n{code_text}\n```\n\n"

    def append_image(img_tag):
        nonlocal content_text
        img_url = img_tag.get('src') or img_tag.get('data-src')
        if img_url:
            local_image = download_image(img_url, outdir)
            if local_image:
                content_text += f'[image: {local_image}]\n'
            else:
                content_text += f'[image: {img_url} (download failed)]\n'
            content_text += '\n'

    def append_list(list_tag, ordered=False):
        nonlocal content_text
        index = 1
        for li in list_tag.find_all('li', recursive=False):
            item_text = li.get_text(strip=True)
            if ordered:
                content_text += f"{index}. {item_text}\n"
                index += 1
            else:
                content_text += f"* {item_text}\n"
        content_text += '\n'

    def walk(node):
        nonlocal content_text
        if isinstance(node, NavigableString):
            return
        if not hasattr(node, 'name'):
            return
        if node.name in {'script', 'style', 'noscript'}:
            return

        if node.name == 'h1':
            content_text += f"# {node.get_text(strip=True)}\n\n"
            return
        if node.name == 'h2':
            content_text += f"## {node.get_text(strip=True)}\n\n"
            return
        if node.name == 'p':
            text = node.get_text(strip=True)
            if text:
                content_text += text + '\n\n'
            return
        if node.name in {'pre', 'code'}:
            if node.name == 'code' and node.parent and node.parent.name == 'pre':
                return
            append_code_block(node.get_text())
            return
        if node.name == 'div':
            classes = node.get('class', [])
            if 'wp-block-code' in classes or 'wp-block-preformatted' in classes:
                pre = node.find('pre')
                code = node.find('code') if pre is None else None
                if pre:
                    append_code_block(pre.get_text())
                elif code:
                    append_code_block(code.get_text())
                return
        if node.name == 'ul':
            append_list(node, ordered=False)
            return
        if node.name == 'ol':
            append_list(node, ordered=True)
            return
        if node.name == 'figure':
            classes = node.get('class', [])
            if 'wp-block-image' in classes or node.find('img'):
                img_tag = node.find('img')
                if img_tag:
                    append_image(img_tag)
                return
        if node.name == 'img':
            append_image(node)
            return

        classes = node.get('class', [])
        if classes and 'social-sharing' in classes:
            return

        for child in node.children:
            walk(child)

    for child in container.children:
        walk(child)

    return content_text


def create_word_document(post_title, publish_date, content_text, outdir, post_url):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    def ensure_code_style(doc):
        try:
            return doc.styles['CodeBlock']
        except KeyError:
            style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Consolas'
            style.font.size = Pt(10)
            style.paragraph_format.left_indent = Inches(0.25)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            return style

    def add_code_paragraph(doc, lines):
        if not lines:
            return
        code_style = ensure_code_style(doc)
        paragraph = doc.add_paragraph(style=code_style)
        for i, line in enumerate(lines):
            if i:
                paragraph.add_run().add_break()
            paragraph.add_run(line)

    doc = Document()
    doc.add_heading(post_title, level=0)
    doc.add_paragraph(publish_date)

    in_code_block = False
    code_lines = []
    for line in content_text.split('\n'):
        if line.strip() == '```':
            if in_code_block:
                add_code_paragraph(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        if '[image:' in line:
            image_path = line[line.find('[image:') + 7:].rstrip(']').strip()
            full_image_path = os.path.join(outdir, image_path)
            if os.path.exists(full_image_path) and os.path.getsize(full_image_path) > 0:
                try:
                    doc.add_picture(full_image_path, width=Pt(350))
                except Exception as e:
                    log.error(f"Failed to embed image {full_image_path} in document: {e}")
                    doc.add_paragraph(f"[Error embedding image: {image_path}]")
        elif line.strip():
            doc.add_paragraph(line)

    if in_code_block and code_lines:
        add_code_paragraph(doc, code_lines)

    # Insert post link as a marker paragraph, followed by the URL (moved to end)
    doc.add_paragraph("Original post link")
    doc.add_paragraph(post_url)

    return doc

def sanitize_filename(title):
    '''
    Sanitizes the title to be safe for use as a filename.
    '''
    # Replace all spaces with underscores
    log.debug(f'+++ Sanitizing title: {title}')
    t1 = title.replace(' ', '_')
    t2 = t1.replace("’", "")
    san_title = re.sub(r'[^\w_]', '-', t2)
    log.debug(f'san_title: {san_title}')
    return san_title


# Helper: check for duplicate filenames
def is_duplicate_filename(text_filename, docx_filename, existing_txt_files, existing_docx_files):
    """
    Returns True if either the text or docx filename already exists in the respective sets.
    """
    return text_filename in existing_txt_files or docx_filename in existing_docx_files

# Helper: save text file
def save_text_file(post_title, publication_date, content_text, post_url, outdir, filename_with_date, extension='txt'):
    """
    Save the blog post as a Markdown text file.
    """
    md_text = f"# {post_title}\nDate: {publication_date}\n\n{content_text}"
    ext = extension.lstrip('.')
    text_filepath = os.path.join(outdir, f"{filename_with_date}.{ext}")
    with open(text_filepath, 'w') as file:
        file.write(md_text + f"\nOriginal post: {post_url}\n")
    log.info(f'Article saved as: {text_filepath}')
    return text_filepath

# Helper: save Word file
def save_word_file(post_title, publication_date, content_text, outdir, filename_with_date, word, log, post_url):
    """
    Save the blog post as a Word document, if requested.
    """
    if word:
        doc_title = post_title.replace('_', ' ').title()
        doc = create_word_document(doc_title, publication_date, content_text, outdir, post_url)
        word_filepath = os.path.join(outdir, filename_with_date + '.docx')
        os.makedirs(os.path.dirname(word_filepath), exist_ok=True)
        doc.save(word_filepath)
        log.info(f'Article saved as: {word_filepath}')
        return word_filepath
    return None


def save_meta_file(post_data, outdir, filename_with_date):
    meta_path = os.path.join(outdir, f"{filename_with_date}.meta.json")
    with open(meta_path, 'w') as file:
        import json
        json.dump(post_data, file, indent=2)
    log.info(f'Metadata saved as: {meta_path}')
    return meta_path


def _select_columns(rows):
    if not rows:
        return []
    keys = set().union(*(row.keys() for row in rows))
    ops = {row.get('operation') for row in rows}
    download_keys = {'text_path', 'md_path', 'docx_path', 'meta_path'}
    if len(ops) == 1:
        op = next(iter(ops))
        op_map = {
            'get-plugins': ['name', 'operation', 'plugin', 'status', 'version'],
            'get-posts': ['title', 'date', 'url', 'txt', 'md', 'docx', 'meta'],
            'list-posts': ['id', 'title', 'status', 'date', 'link'],
            'get-pages': ['id', 'title', 'status', 'date', 'link'],
            'get-categories': ['id', 'name', 'slug', 'count', 'parent', 'description'],
            'get-tags': ['id', 'name', 'slug', 'count', 'description'],
            'get-users': ['id', 'name', 'registered_date', 'roles', 'capabilities'],
            'get-user-me': ['id', 'name', 'registered_date', 'roles', 'capabilities'],
            'get-media': ['id', 'title', 'media_type', 'mime_type', 'link'],
            'get-comments': ['id', 'post', 'author_name', 'status', 'date'],
            'get-types': ['name', 'slug', 'rest_base', 'description'],
            'get-statuses': ['name', 'label'],
            'get-taxonomies': ['name', 'slug', 'description'],
            'get-settings': ['key', 'value'],
            'get-themes': ['stylesheet', 'name', 'version', 'status'],
            'shift-scheduled': ['id', 'title', 'from', 'to'],
            'backfill-post-navigation': ['id', 'title', 'status', 'date', 'action', 'dry_run', 'link'],
        }
        if op in op_map:
            return op_map[op]
    if keys.issuperset({'name', 'operation', 'plugin', 'status', 'version'}):
        return ['name', 'operation', 'plugin', 'status', 'version']
    if keys.issuperset({'title', 'date', 'url'}) and keys.intersection(download_keys):
        return ['title', 'date', 'url', 'txt', 'md', 'docx', 'meta']
    return sorted(keys)


def render_ascii_table(rows, max_width=TABLE_MAX_WIDTH):
    if not rows:
        return ''
    columns = _select_columns(rows)
    if not columns:
        return ''
    def _format_value(col, value, row):
        if value is None:
            value = ''
        if col in {'txt', 'md', 'docx', 'meta'}:
            key_map = {
                'txt': 'text_path',
                'md': 'md_path',
                'docx': 'docx_path',
                'meta': 'meta_path',
            }
            value = 'X' if row.get(key_map[col]) else ''
            return value
        value = str(value)
        if col == 'url' and (row.get('operation') == 'get-posts' or any(k in row for k in ('text_path', 'md_path', 'docx_path', 'meta_path'))):
            parsed = urlparse(value)
            slug = parsed.path.strip('/')
            return slug or value
        if col == 'registered_date' and 'T' in value:
            value = value.split('T', 1)[0]
        return value

    widths = {col: len(col) for col in columns}
    for row in rows:
        for col in columns:
            val = _format_value(col, row.get(col, ''), row)
            widths[col] = max(widths[col], len(val))

    def _truncate(value, width):
        if len(value) <= width:
            return value
        if width <= 3:
            return value[:width]
        return value[:width - 3] + '...'

    def _format_row(values, header=False):
        parts = []
        for col in columns:
            if header:
                raw = col
            else:
                raw = _format_value(col, values.get(col, '') if values else '', values or {})
            parts.append(_truncate(raw, widths[col]).ljust(widths[col]))
        return "| " + " | ".join(parts) + " |"

    def _line_length():
        return len(_format_row({col: col for col in columns}))

    min_widths = {col: TABLE_MIN_WIDTH for col in columns}
    if 'title' in min_widths:
        min_widths['title'] = TABLE_MIN_WIDTH_TITLE
    if 'date' in min_widths:
        min_widths['date'] = TABLE_MIN_WIDTH_DATE
    while _line_length() > max_width:
        shrinkable = [c for c in columns if widths[c] > min_widths.get(c, 4)]
        if not shrinkable:
            break
        widest = max(shrinkable, key=lambda c: widths[c])
        widths[widest] -= 1

    border = "+-" + "-+-".join("-" * widths[col] for col in columns) + "-+"
    header = _format_row({}, header=True)
    lines = [border, header, border]
    for row in rows:
        lines.append(_format_row(row))
    lines.append(border)
    return "\n".join(lines)


def render_results(rows, outfile_format):
    if not rows:
        return ''
    fmt = (outfile_format or 'csv').lower()
    if fmt not in {'csv', 'json'}:
        raise ValueError(f"Invalid outfile format: {outfile_format}")
    if fmt == 'json':
        import json
        return json.dumps(rows, indent=2)
    fieldnames = sorted({k for row in rows for k in row.keys()})
    lines = [','.join(fieldnames)]
    for row in rows:
        values = []
        for key in fieldnames:
            value = row.get(key, "")
            if value is None:
                value = ""
            value = str(value).replace('\n', ' ').replace('\r', ' ')
            if ',' in value or '"' in value:
                value = '"' + value.replace('"', '""') + '"'
            values.append(value)
        lines.append(','.join(values))
    return '\n'.join(lines)


def write_results(outfile_base, outfile_format, rows):
    if not rows:
        return None
    fmt = (outfile_format or 'csv').lower()
    outfile_base = outfile_base or 'out'
    if fmt == 'json':
        path = outfile_base if outfile_base.endswith('.json') else f"{outfile_base}.json"
    else:
        path = outfile_base if outfile_base.endswith('.csv') else f"{outfile_base}.csv"
    content = render_results(rows, outfile_format)
    with open(path, 'w') as f:
        f.write(content + '\n')
    return path


# Helper function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.clear_content()
    paragraph._p.append(hyperlink)

def build_wp_api_base(url):
    url = (url or '').strip()
    if not url:
        return ''

    if '://' not in url:
        url = f"https://{url}"

    normalized = url.rstrip('/')
    if '/wp-json/' in normalized:
        base = normalized.split('/wp-json/')[0]
        return base.rstrip('/')

    parsed = urlparse(normalized)
    if parsed.scheme and parsed.netloc:
        return f"{parsed.scheme}://{parsed.netloc}"

    return normalized

def build_wp_api_posts_url(url):
    url = (url or '').strip()
    if not url:
        return ''
    normalized = url.rstrip('/')

    if '/wp-json/' in normalized:
        if normalized.endswith('/wp/v2/posts'):
            return normalized
        if normalized.endswith('/wp-json/wp/v2'):
            return normalized + '/posts'
        if normalized.endswith('/wp-json'):
            return normalized + '/wp/v2/posts'

    base = build_wp_api_base(normalized)
    if not base:
        return ''
    return base.rstrip('/') + '/wp-json/wp/v2/posts'

def build_wp_api_plugins_url(url):
    base = build_wp_api_base(url)
    if not base:
        return ''
    return base.rstrip('/') + '/wp-json/wp/v2/plugins'

def build_auth_header(username, app_password):
    token = base64.b64encode(
        f"{username}:{app_password}".encode()).decode("utf-8")
    return {"Authorization": f"Basic {token}"}

def fetch_wp_plugins(url, headers):
    plugins_url = build_wp_api_plugins_url(url)
    if not plugins_url:
        log.error("Invalid URL for WordPress API.")
        return None
    plugins = []
    page = 1
    while True:
        params = {'per_page': WP_API_MAX_PER_PAGE, 'page': page}
        response = requests.get(plugins_url, headers=headers, params=params)
        if response.status_code >= 400:
            log.error(f"Failed to fetch plugins: {response.status_code} {response.text}")
            return None
        data = response.json()
        plugins.extend(data)
        total_pages = int(response.headers.get('X-WP-TotalPages', 1))
        if page >= total_pages:
            break
        page += 1
    return plugins


def fetch_wp_endpoint(url, endpoint, headers, params=None):
    base = build_wp_api_base(url)
    if not base:
        log.error("Invalid URL for WordPress API.")
        return None
    endpoint = endpoint.lstrip('/')
    full_url = f"{base.rstrip('/')}/wp-json/wp/v2/{endpoint}"
    params = dict(params or {})
    params.setdefault('per_page', WP_API_MAX_PER_PAGE)
    page = 1
    items = []
    while True:
        params['page'] = page
        response = requests.get(full_url, headers=headers, params=params)
        if response.status_code >= 400:
            log.error(f"Failed to fetch {endpoint}: {response.status_code} {response.text}")
            return None
        data = response.json()
        if isinstance(data, list):
            items.extend(data)
            total_pages = int(response.headers.get('X-WP-TotalPages', 1))
            if page >= total_pages:
                break
            page += 1
        else:
            return data
    return items


def normalize_wp_rows(operation, data):
    if data is None:
        return []
    rows = []
    if isinstance(data, dict):
        if operation == 'get-settings':
            return [{'operation': operation, 'key': k, 'value': v} for k, v in data.items()]
        if operation in {'get-types', 'get-statuses', 'get-taxonomies'}:
            for key, value in data.items():
                if isinstance(value, dict):
                    row = {'operation': operation, 'name': key}
                    row.update({
                        'slug': value.get('slug') or '',
                        'rest_base': value.get('rest_base') or '',
                        'description': value.get('description') or '',
                        'label': value.get('label') or ''
                    })
                    rows.append(row)
            return rows
        if operation == 'get-user-me':
            roles = data.get('roles') or []
            capabilities = data.get('capabilities') or {}
            if isinstance(capabilities, dict):
                cap_list = [k for k, v in capabilities.items() if v]
            else:
                cap_list = []
            return [{
                'operation': operation,
                'id': data.get('id'),
                'name': data.get('name') or '',
                'slug': data.get('slug') or '',
                'link': data.get('link') or '',
                'registered_date': data.get('registered_date') or '',
                'roles': ';'.join(roles) if isinstance(roles, list) else str(roles),
                'capabilities': ';'.join(cap_list)
            }]
        return [{'operation': operation, **data}]

    def _title(item):
        title_html = (item.get('title') or {}).get('rendered', '')
        return BeautifulSoup(title_html, 'html.parser').get_text().strip()

    for item in data:
        if operation in {'list-posts', 'get-pages'}:
            rows.append({
                'operation': operation,
                'id': item.get('id'),
                'title': _title(item),
                'status': item.get('status') or '',
                'date': item.get('date') or '',
                'link': item.get('link') or ''
            })
        elif operation in {'get-categories', 'get-tags'}:
            rows.append({
                'operation': operation,
                'id': item.get('id'),
                'name': item.get('name') or '',
                'slug': item.get('slug') or '',
                'count': item.get('count') or 0
            })
        elif operation == 'get-users':
            roles = item.get('roles') or []
            capabilities = item.get('capabilities') or {}
            if isinstance(capabilities, dict):
                cap_list = [k for k, v in capabilities.items() if v]
            else:
                cap_list = []
            rows.append({
                'operation': operation,
                'id': item.get('id'),
                'name': item.get('name') or '',
                'slug': item.get('slug') or '',
                'link': item.get('link') or '',
                'registered_date': item.get('registered_date') or '',
                'roles': ';'.join(roles) if isinstance(roles, list) else str(roles),
                'capabilities': ';'.join(cap_list)
            })
        elif operation == 'get-media':
            title_html = (item.get('title') or {}).get('rendered', '')
            rows.append({
                'operation': operation,
                'id': item.get('id'),
                'title': BeautifulSoup(title_html, 'html.parser').get_text().strip(),
                'media_type': item.get('media_type') or '',
                'mime_type': item.get('mime_type') or '',
                'link': item.get('link') or ''
            })
        elif operation == 'get-comments':
            rows.append({
                'operation': operation,
                'id': item.get('id'),
                'post': item.get('post') or '',
                'author_name': item.get('author_name') or '',
                'status': item.get('status') or '',
                'date': item.get('date') or ''
            })
        elif operation == 'get-themes':
            rows.append({
                'operation': operation,
                'stylesheet': item.get('stylesheet') or '',
                'name': item.get('name') or '',
                'version': item.get('version') or '',
                'status': item.get('status') or ''
            })
        else:
            rows.append({'operation': operation, **item})
    return rows

def format_wp_api_date(date_str):
    if not date_str:
        return "unknown-date"
    try:
        dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        return dt.strftime('%Y-%m-%d')
    except ValueError as e:
        log.error(f"Error parsing API date '{date_str}': {e}")
    return "unknown-date"

def parse_date_range(date_arg):
    if not date_arg:
        return None, None
    today = datetime.today().date()
    arg = date_arg.strip().lower()
    if arg == 'all':
        return None, None
    if arg == 'today':
        start = end = today
    elif arg == 'week':
        start = today - timedelta(days=7)
        end = today
    elif arg == 'month':
        start = today - timedelta(days=30)
        end = today
    elif arg == 'year':
        start = today - timedelta(days=365)
        end = today
    else:
        try:
            start_str, end_str = date_arg.split(':', 1)
            start = datetime.strptime(start_str, "%m-%d-%Y").date()
            end = datetime.strptime(end_str, "%m-%d-%Y").date()
        except ValueError as exc:
            raise ValueError(f"Invalid --date value: {date_arg}") from exc
    if start > end:
        start, end = end, start
    return start, end

def add_date_range_params(params, date_range):
    if not date_range or not date_range[0] or not date_range[1]:
        return params
    start, end = date_range
    params = dict(params)
    params['after'] = start.strftime('%Y-%m-%dT00:00:00')
    params['before'] = end.strftime('%Y-%m-%dT23:59:59')
    return params


def load_dotenv(path: str = '.env', required: bool = True):
    env_path = os.path.join(os.getcwd(), path)
    if not os.path.exists(env_path):
        if required:
            log.error(f'XXX Missing required env file: {env_path}')
            sys.exit(1)
        return
    try:
        with open(env_path, 'r') as file:
            for line in file:
                line = line.strip()
                if not line or line.startswith('#'):
                    continue
                if '=' not in line:
                    continue
                key, value = line.split('=', 1)
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key and key not in os.environ:
                    os.environ[key] = value
    except Exception as exc:
        log.error(f'XXX Failed to read env file {env_path}: {exc}')
        sys.exit(1)


def parse_export_site_arg(arg_value):
    if not arg_value:
        return set(), False
    raw = [part.strip().lower() for part in arg_value.split(',') if part.strip()]
    if not raw:
        raise ValueError('Invalid --export-site value (empty).')
    if 'all' in raw and 'all-no-media' in raw:
        raise ValueError('Cannot combine all and all-no-media in --export-site.')
    if 'all' in raw:
        return set(EXPORT_RESOURCES), True
    if 'all-no-media' in raw:
        return set(EXPORT_RESOURCES), False
    unknown = [item for item in raw if item not in EXPORT_RESOURCES]
    if unknown:
        raise ValueError(f'Unknown --export-site option(s): {", ".join(sorted(unknown))}')
    return set(raw), ('media' in raw)


def fetch_wp_root(url, headers):
    base = build_wp_api_base(url)
    if not base:
        log.error("Invalid URL for WordPress API.")
        return None
    root_url = f"{base.rstrip('/')}/wp-json"
    response = requests.get(root_url, headers=headers)
    if response.status_code >= 400:
        log.error(f"Failed to fetch WP root: {response.status_code} {response.text}")
        return None
    return response.json()


def fetch_wp_menus(url, headers):
    base = build_wp_api_base(url)
    if not base:
        log.error("Invalid URL for WordPress API.")
        return None
    candidates = [
        f"{base.rstrip('/')}/wp-json/wp/v2/menus",
        f"{base.rstrip('/')}/wp-json/wp-api-menus/v2/menus"
    ]
    for menu_url in candidates:
        response = requests.get(menu_url, headers=headers, params={'per_page': WP_API_MAX_PER_PAGE})
        if response.status_code == 404:
            continue
        if response.status_code >= 400:
            log.error(f"Failed to fetch menus: {response.status_code} {response.text}")
            return None
        data = response.json()
        return data
    log.warning("Menus endpoint not available via REST API.")
    return None


def _load_existing_ids(path):
    ids = set()
    if not os.path.exists(path):
        return ids
    with open(path, 'r') as file:
        for line in file:
            line = line.strip()
            if not line:
                continue
            try:
                item = json.loads(line)
            except json.JSONDecodeError:
                continue
            item_id = item.get('id')
            if item_id is not None:
                ids.add(item_id)
    return ids


def _write_json_file(path, data):
    with open(path, 'w') as file:
        json.dump(data, file, indent=2)


def _write_jsonl(path, items, incremental=False, existing_ids=None):
    if incremental:
        if existing_ids is None:
            existing_ids = _load_existing_ids(path)
        mode = 'a'
    else:
        existing_ids = set()
        mode = 'w'
    written = 0
    skipped = 0
    with open(path, mode) as file:
        for item in items:
            item_id = item.get('id')
            if incremental and item_id is not None and item_id in existing_ids:
                skipped += 1
                continue
            file.write(json.dumps(item, ensure_ascii=True) + '\n')
            written += 1
    return written, skipped, len(items)


def _download_media_files(media_items, outdir, incremental=False, existing_ids=None):
    media_dir = os.path.join(outdir, 'media')
    os.makedirs(media_dir, exist_ok=True)
    downloaded = 0
    skipped_existing = 0
    skipped_missing = 0
    failed = 0
    if incremental and existing_ids is None:
        existing_ids = set()
    for item in media_items:
        item_id = item.get('id')
        if incremental and item_id is not None and item_id in existing_ids:
            skipped_existing += 1
            continue
        source_url = item.get('source_url') or ''
        if not source_url:
            skipped_missing += 1
            continue
        filename = os.path.basename(urlparse(source_url).path)
        if not filename:
            filename = f"media_{item_id or 'unknown'}"
        filename = f"{item_id}_{valid_filename(filename)}" if item_id is not None else valid_filename(filename)
        dest_path = os.path.join(media_dir, filename)
        if incremental and os.path.exists(dest_path):
            skipped_existing += 1
            continue
        try:
            response = requests.get(source_url, stream=True)
            if response.status_code >= 400:
                failed += 1
                continue
            with open(dest_path, 'wb') as file:
                for chunk in response.iter_content(8192):
                    file.write(chunk)
            item['exported_path'] = dest_path
            downloaded += 1
        except requests.RequestException:
            failed += 1
    return {
        'downloaded': downloaded,
        'skipped_existing': skipped_existing,
        'skipped_missing_url': skipped_missing,
        'failed': failed
    }


def export_site(url, headers, export_items, outdir, incremental, download_media):
    export_summary = {
        'resources': {},
        'warnings': []
    }
    os.makedirs(outdir, exist_ok=True)

    log.info(f'+  Exporting site data to: {outdir}')
    root_info = fetch_wp_root(url, headers)
    if root_info is None:
        export_summary['warnings'].append('Failed to fetch site root info')
    else:
        root_path = os.path.join(outdir, 'site.json')
        _write_json_file(root_path, root_info)
        export_summary['resources']['site'] = {
            'path': root_path,
            'count': 1
        }
        log.info(f'+  Wrote site.json: {root_path}')

    resource_map = {
        'posts': ('posts', 'posts.jsonl'),
        'pages': ('pages', 'pages.jsonl'),
        'media': ('media', 'media.jsonl'),
        'comments': ('comments', 'comments.jsonl'),
        'users': ('users', 'users.jsonl'),
        'categories': ('categories', 'categories.jsonl'),
        'tags': ('tags', 'tags.jsonl'),
        'taxonomies': ('taxonomies', 'taxonomies.json'),
        'types': ('types', 'types.json'),
        'statuses': ('statuses', 'statuses.json'),
        'settings': ('settings', 'settings.json'),
        'menus': ('menus', 'menus.jsonl'),
        'plugins': ('plugins', 'plugins.jsonl')
    }
    for resource in sorted(export_items):
        endpoint, filename = resource_map[resource]
        params = {}
        if resource in {'posts', 'pages', 'media', 'users'} and headers.get('Authorization'):
            params['context'] = 'edit'
        if resource in {'posts', 'pages'} and headers.get('Authorization'):
            params['status'] = 'any'
        log.info(f'+  Exporting {resource} via WP API')
        if resource == 'menus':
            data = fetch_wp_menus(url, headers)
        elif resource == 'plugins':
            if not headers.get('Authorization'):
                export_summary['warnings'].append('Plugins export skipped: missing credentials')
                log.info('+  Plugins export skipped: missing credentials')
                continue
            data = fetch_wp_plugins(url, headers)
        else:
            data = fetch_wp_endpoint(url, endpoint, headers, params=params)
        if data is None:
            export_summary['warnings'].append(f'Failed to fetch {resource}')
            log.info(f'+  Export failed for {resource}')
            continue
        path = os.path.join(outdir, filename)
        if isinstance(data, list):
            existing_ids = _load_existing_ids(path) if incremental else None
            log.info(f'+  {resource} fetched: {len(data)} items')
            if resource == 'media' and download_media:
                log.info('+  Downloading media binaries')
                media_stats = _download_media_files(data, outdir, incremental, existing_ids)
                export_summary['resources']['media_files'] = {
                    'downloaded': media_stats['downloaded'],
                    'skipped_existing': media_stats['skipped_existing'],
                    'skipped_missing_url': media_stats['skipped_missing_url'],
                    'failed': media_stats['failed'],
                    'path': os.path.join(outdir, 'media')
                }
                log.info(
                    f'+  Media files: downloaded={media_stats["downloaded"]}, '
                    f'skipped_existing={media_stats["skipped_existing"]}, '
                    f'skipped_missing_url={media_stats["skipped_missing_url"]}, '
                    f'failed={media_stats["failed"]}'
                )
            written, skipped, total = _write_jsonl(path, data, incremental=incremental, existing_ids=existing_ids)
            export_summary['resources'][resource] = {
                'path': path,
                'count': total,
                'written': written,
                'skipped_existing': skipped
            }
            log.info(f'+  Wrote {resource}: {path} (written={written}, skipped_existing={skipped})')
        else:
            _write_json_file(path, data)
            export_summary['resources'][resource] = {
                'path': path,
                'count': len(data) if isinstance(data, dict) else 1
            }
            log.info(f'+  Wrote {resource}: {path}')

    manifest_path = os.path.join(outdir, 'manifest.json')
    manifest = {
        'generated_at': datetime.utcnow().isoformat() + 'Z',
        'source_url': build_wp_api_base(url),
        'incremental': bool(incremental),
        'download_media': bool(download_media),
        'resources': export_summary['resources'],
        'warnings': export_summary['warnings']
    }
    _write_json_file(manifest_path, manifest)
    export_summary['manifest'] = manifest_path
    log.info(f'+  Wrote manifest: {manifest_path}')
    return export_summary


def read_markdown_content(path):
    if path == '-':
        return sys.stdin.read()
    with open(path, 'r') as file:
        return file.read()


def load_meta_json(path):
    with open(path, 'r') as file:
        data = json.load(file)
    if not isinstance(data, dict):
        raise ValueError('Metadata JSON must be an object.')
    return data


def strip_leading_h1(markdown_text):
    lines = markdown_text.splitlines()
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    if idx < len(lines) and re.match(r'^\s*#\s+\S', lines[idx]):
        del lines[idx]
        if idx < len(lines) and not lines[idx].strip():
            del lines[idx]
    return "\n".join(lines)


def resolve_term_ids(url, headers, endpoint, names, create_missing=False, warn_missing=False):
    ids = []
    for name in names:
        params = {'search': name}
        data = fetch_wp_endpoint(url, endpoint, headers, params=params)
        if not isinstance(data, list):
            raise ValueError(f'Failed to fetch {endpoint} for {name}')
        matches = []
        for item in data:
            if (item.get('name') or '').lower() == name.lower():
                matches.append(item)
        if not matches:
            if not create_missing:
                if warn_missing:
                    log.warning(f"Missing {endpoint} name '{name}', skipping")
                    continue
                raise ValueError(f'No {endpoint} found for name: {name}')
            term = create_wp_term(url, headers, endpoint, name)
            ids.append(term.get('id'))
        else:
            ids.append(matches[0].get('id'))
    return ids


def create_wp_term(url, headers, endpoint, name):
    base = build_wp_api_base(url)
    if not base:
        raise ValueError("Invalid URL for WordPress API.")
    if not headers.get('Authorization'):
        raise ValueError(f"Missing credentials to create {endpoint}: {name}")
    endpoint = endpoint.lstrip('/')
    full_url = f"{base.rstrip('/')}/wp-json/wp/v2/{endpoint}"
    response = requests.post(full_url, headers=headers, json={'name': name})
    if response.status_code >= 400:
        raise ValueError(f"Failed to create {endpoint} '{name}': {response.status_code} {response.text}")
    return response.json()


def upload_media_and_replace(markdown_text, base_dir, url, headers):
    image_pattern = re.compile(r'!\[[^\]]*]\(([^)\s]+)(?:\s+"[^"]*")?\)')
    matches = list(image_pattern.finditer(markdown_text))
    if not matches:
        return markdown_text, []
    uploads = []
    updated = markdown_text
    media_url = f"{build_wp_api_base(url).rstrip('/')}/wp-json/wp/v2/media"
    for match in matches:
        raw_path = match.group(1)
        if raw_path.startswith('http://') or raw_path.startswith('https://') or raw_path.startswith('data:'):
            continue
        local_path = raw_path
        if not os.path.isabs(local_path):
            local_path = os.path.join(base_dir, local_path)
        if not os.path.exists(local_path):
            raise ValueError(f"Image file not found: {local_path}")
        filename = os.path.basename(local_path)
        mime_type, _ = mimetypes.guess_type(local_path)
        mime_type = mime_type or 'application/octet-stream'
        with open(local_path, 'rb') as file:
            response = requests.post(
                media_url,
                headers=headers,
                files={'file': (filename, file, mime_type)}
            )
        if response.status_code >= 400:
            raise ValueError(f"Media upload failed: {response.status_code} {response.text}")
        media_item = response.json()
        source_url = media_item.get('source_url') or ''
        if not source_url:
            raise ValueError('Media upload succeeded but no source_url returned.')
        updated = updated.replace(raw_path, source_url)
        uploads.append({'id': media_item.get('id'), 'source_url': source_url, 'file': local_path})
    return updated, uploads


def find_local_markdown_images(markdown_text, base_dir):
    image_pattern = re.compile(r'!\[[^\]]*]\(([^)\s]+)(?:\s+"[^"]*")?\)')
    local_paths = []
    for match in image_pattern.finditer(markdown_text):
        raw_path = match.group(1)
        if raw_path.startswith('http://') or raw_path.startswith('https://') or raw_path.startswith('data:'):
            continue
        local_path = raw_path
        if not os.path.isabs(local_path):
            local_path = os.path.join(base_dir, local_path)
        local_paths.append(local_path)
    return local_paths


def fetch_latest_scheduled_date(url, headers, tz_name):
    posts_url = build_wp_api_posts_url(url)
    params = {
        'status': 'future',
        'per_page': 1,
        'orderby': 'date',
        'order': 'desc'
    }
    response = requests.get(posts_url, headers=headers, params=params)
    if response.status_code >= 400:
        raise ValueError(f"Failed to fetch scheduled posts: {response.status_code} {response.text}")
    data = response.json()
    if not isinstance(data, list) or not data:
        return None
    latest = data[0]
    date_str = latest.get('date')
    if not date_str:
        return None
    dt = datetime.fromisoformat(date_str)
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=ZoneInfo(tz_name))
    return dt


def parse_publish_date(date_str: str) -> date:
    if not date_str:
        raise ValueError('publish-date is required')
    try:
        return datetime.strptime(date_str, '%m-%d-%Y').date()
    except ValueError as exc:
        raise ValueError('publish-date must be MM-DD-YYYY') from exc


def _coerce_publish_date(publish_date: Any) -> date | None:
    if not publish_date:
        return None
    if isinstance(publish_date, datetime):
        return publish_date.date()
    if isinstance(publish_date, date):
        return publish_date
    if isinstance(publish_date, str):
        return parse_publish_date(publish_date)
    raise ValueError('publish-date must be MM-DD-YYYY')


def fetch_scheduled_posts(url, headers, tz_name, start_date: date | None = None, exclude_ids: set[int] | None = None):
    posts_url = build_wp_api_posts_url(url)
    params = {
        'status': 'future',
        'per_page': WP_API_MAX_PER_PAGE,
        'orderby': 'date',
        'order': 'asc',
        'page': 1,
    }
    posts: list[dict[str, Any]] = []
    while True:
        response = requests.get(posts_url, headers=headers, params=params)
        if response.status_code >= 400:
            raise ValueError(f"Failed to fetch scheduled posts: {response.status_code} {response.text}")
        data = response.json()
        if not isinstance(data, list) or not data:
            break
        for post in data:
            post_id = post.get('id')
            if exclude_ids and post_id in exclude_ids:
                continue
            date_str = post.get('date')
            if not date_str:
                continue
            dt = datetime.fromisoformat(date_str)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=ZoneInfo(tz_name))
            if start_date and dt.date() < start_date:
                continue
            title_raw = ''
            title_obj = post.get('title') or {}
            if isinstance(title_obj, dict):
                title_raw = title_obj.get('rendered', '') or ''
            title_text = BeautifulSoup(title_raw, 'html.parser').get_text().strip() if title_raw else ''
            posts.append({'id': post_id, 'date': dt, 'title': title_text})
        total_pages = int(response.headers.get('X-WP-TotalPages', 1))
        if params['page'] >= total_pages:
            break
        params['page'] += 1
    return posts


def shift_scheduled_posts(url, headers, start_date: date, tz_name=DEFAULT_TIMEZONE, dry_run=False, exclude_ids: set[int] | None = None):
    if not start_date:
        return []
    posts = fetch_scheduled_posts(url, headers, tz_name, start_date=start_date, exclude_ids=exclude_ids)
    if not posts:
        return []
    posts_url = build_wp_api_posts_url(url)
    shifted = []
    posts.sort(key=lambda p: p.get('date'), reverse=True)
    for post in posts:
        post_id = post.get('id')
        if not post_id:
            continue
        old_dt = post['date']
        new_dt = old_dt + timedelta(days=1)
        shifted.append({'id': post_id, 'title': post.get('title', ''), 'from': old_dt, 'to': new_dt})
        if dry_run:
            continue
        update_url = posts_url.rstrip('/') + f'/{post_id}'
        payload = {
            'date': new_dt.strftime('%Y-%m-%dT%H:%M:%S'),
            'status': 'future'
        }
        response = requests.post(update_url, headers=headers, json=payload)
        if response.status_code >= 400:
            raise ValueError(f"Failed to shift post {post_id}: {response.status_code} {response.text}")
    return shifted


def format_shift_report(shifted: list[dict[str, Any]]):
    report = []
    for item in shifted or []:
        report.append({
            'operation': 'shift-scheduled',
            'id': item.get('id'),
            'title': item.get('title', ''),
            'from': item.get('from').strftime('%Y-%m-%dT%H:%M:%S') if item.get('from') else '',
            'to': item.get('to').strftime('%Y-%m-%dT%H:%M:%S') if item.get('to') else '',
        })
    return report


def has_post_navigation_blocks(content: str) -> bool:
    return 'wp:post-navigation-link' in (content or '')


def split_post_navigation_blocks(content: str) -> tuple[str, str]:
    match = POST_NAVIGATION_BLOCKS_PATTERN.search(content or '')
    if not match:
        return (content or ''), ''
    body = (content or '')[:match.start()].rstrip()
    nav = match.group(0)
    return body, nav


def detect_post_body_format(content: str) -> str:
    body = (content or '').strip()
    if not body:
        return 'empty'
    if '<!-- wp:' in body:
        return 'gutenberg'
    if re.search(r'<(?:p|ul|ol|li|h[1-6]|blockquote|pre|figure|div|table|hr)\b', body, re.I):
        return 'html'
    return 'markdown'


def ensure_post_navigation_blocks(content: str) -> str:
    if has_post_navigation_blocks(content):
        return content
    body = (content or '').rstrip()
    if not body:
        return POST_NAVIGATION_BLOCKS
    return f'{body}\n\n{POST_NAVIGATION_BLOCKS}'


def normalize_post_navigation_content(content: str) -> tuple[str, str]:
    body, nav = split_post_navigation_blocks(content)
    body_format = detect_post_body_format(body)
    normalized_body = body.rstrip()

    if body_format == 'markdown':
        normalized_body = render_markdown((body or '').strip())

    if nav:
        if not normalized_body:
            return nav, body_format
        return f'{normalized_body}\n\n{nav}', body_format

    return ensure_post_navigation_blocks(normalized_body), body_format


def fetch_wp_post_edit(url, headers, post_id):
    posts_url = build_wp_api_posts_url(url)
    if not posts_url:
        raise ValueError('Invalid URL for WordPress API.')
    response = requests.get(
        posts_url.rstrip('/') + f'/{post_id}',
        headers=headers,
        params={'context': 'edit'},
    )
    if response.status_code >= 400:
        raise ValueError(f'Failed to fetch post {post_id}: {response.status_code} {response.text}')
    return response.json()


def backfill_post_navigation(
    url,
    headers,
    dry_run=False,
    status='publish',
    limit=None,
    post_id=None,
):
    if post_id is not None:
        posts = [fetch_wp_post_edit(url, headers, post_id)]
    else:
        params = {
            'context': 'edit',
            'status': status,
            '_fields': 'id,slug,date,status,link,title',
        }
        posts = fetch_wp_endpoint(url, 'posts', headers, params=params)
        if not isinstance(posts, list):
            raise ValueError('Failed to fetch posts for backfill.')

    rows = []
    posts_url = build_wp_api_posts_url(url)
    processed = 0
    updated = 0
    already = 0
    skipped = 0

    for post in posts:
        if limit is not None and processed >= limit:
            break
        processed += 1
        post_id = post.get('id')
        detail = post if post_id is not None and 'content' in post else fetch_wp_post_edit(url, headers, post_id)
        title_html = ((post.get('title') or {}).get('rendered', '')) if isinstance(post.get('title'), dict) else ''
        title = BeautifulSoup(title_html, 'html.parser').get_text().strip() if title_html else ''
        if not title:
            detail_title_html = ((detail.get('title') or {}).get('rendered', '')) if isinstance(detail.get('title'), dict) else ''
            title = BeautifulSoup(detail_title_html, 'html.parser').get_text().strip() if detail_title_html else ''
        content_obj = detail.get('content') or {}
        raw_content = content_obj.get('raw') or ''
        row = {
            'operation': 'backfill-post-navigation',
            'id': post_id or '',
            'title': title,
            'status': detail.get('status') or post.get('status') or '',
            'date': format_wp_api_date(detail.get('date') or post.get('date') or ''),
            'link': detail.get('link') or post.get('link') or '',
            'dry_run': str(bool(dry_run)),
        }
        if not raw_content:
            row['action'] = 'skipped-no-raw-content'
            rows.append(row)
            skipped += 1
            continue
        if has_post_navigation_blocks(raw_content):
            row['action'] = 'already-has-navigation'
            rows.append(row)
            already += 1
            continue

        updated_content, body_format = normalize_post_navigation_content(raw_content)
        row['action'] = 'would-update' if dry_run else 'updated'
        row['content_changed'] = str(updated_content != raw_content)
        row['content_format'] = body_format
        rows.append(row)
        if dry_run:
            updated += 1
            continue

        update_url = posts_url.rstrip('/') + f'/{post_id}'
        response = requests.post(update_url, headers=headers, json={'content': updated_content})
        if response.status_code >= 400:
            raise ValueError(
                f'Failed to update post {post_id}: {response.status_code} {response.text}'
            )
        updated += 1

    stats = {
        'processed': processed,
        'already_has_navigation': already,
        'needs_update': updated,
        'skipped': skipped,
    }
    return rows, stats


def build_schedule_payload(url, headers, content_md, meta, tz_name=DEFAULT_TIMEZONE, publish_date: Any | None = None):
    title = meta.get('title') or meta.get('post_title')
    if not title:
        raise ValueError('Metadata JSON must include title.')
    categories = meta.get('categories') or []
    if isinstance(categories, str):
        categories = [categories]
    if not isinstance(categories, list):
        raise ValueError('categories must be a list of category names.')
    categories = [c for c in categories if c]
    if 'The250' not in categories:
        categories.append('The250')
    seen = set()
    deduped = []
    for name in categories:
        key = name.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(name)
    categories = deduped
    required_categories = [name for name in categories if name.lower() == 'the250']
    optional_categories = [name for name in categories if name.lower() != 'the250']
    category_ids = []
    if required_categories:
        category_ids.extend(resolve_term_ids(url, headers, 'categories', required_categories, create_missing=True))
    if optional_categories:
        category_ids.extend(resolve_term_ids(url, headers, 'categories', optional_categories, create_missing=True))
    tags = meta.get('tags') or []
    if isinstance(tags, str):
        tags = [tags]
    if tags:
        if not isinstance(tags, list):
            raise ValueError('tags must be a list of tag names.')
        tag_ids = resolve_term_ids(url, headers, 'tags', tags, create_missing=True)
    else:
        tag_ids = []

    publish_day = _coerce_publish_date(publish_date)
    if publish_day:
        schedule_dt = datetime.combine(publish_day, time(8, 44), tzinfo=ZoneInfo(tz_name))
    else:
        latest_dt = fetch_latest_scheduled_date(url, headers, tz_name)
        now_local = datetime.now(ZoneInfo(tz_name))
        base_date = latest_dt.date() if latest_dt else now_local.date()
        schedule_date = base_date + timedelta(days=1)
        schedule_dt = datetime.combine(schedule_date, time(8, 44), tzinfo=ZoneInfo(tz_name))

    payload = {
        'title': title,
        'content': content_md,
        'status': 'future',
        'date': schedule_dt.strftime('%Y-%m-%dT%H:%M:%S'),
        'categories': category_ids
    }
    if tag_ids:
        payload['tags'] = tag_ids
    if meta.get('excerpt'):
        payload['excerpt'] = meta.get('excerpt')
    slug = (meta.get('slug') or '').strip()
    if not slug:
        slug = re.sub(r'[^a-z0-9]+', '-', title.lower()).strip('-')
    if slug:
        payload['slug'] = slug

    yoast_meta = {}
    focus_keyphrase = (
        meta.get('yoast_wpseo_focuskw')
        or meta.get('yoast_focus_keyphrase')
        or meta.get('focus_keyphrase')
        or ''
    )
    meta_description = (
        meta.get('yoast_wpseo_metadesc')
        or meta.get('yoast_meta_description')
        or meta.get('meta_description')
        or ''
    )
    if focus_keyphrase:
        yoast_meta['yoast_wpseo_focuskw'] = focus_keyphrase
        yoast_meta['_yoast_wpseo_focuskw'] = focus_keyphrase
    if meta_description:
        yoast_meta['yoast_wpseo_metadesc'] = meta_description
        yoast_meta['_yoast_wpseo_metadesc'] = meta_description
    if yoast_meta:
        payload['meta'] = yoast_meta

    return payload, schedule_dt


def find_post_by_slug(url, headers, slug):
    if not slug:
        return None
    base = build_wp_api_base(url)
    if not base:
        log.error("Invalid URL for WordPress API.")
        return None
    params = {
        'slug': slug,
        'per_page': 1,
        'status': 'any',
    }
    if headers.get('Authorization'):
        params['context'] = 'edit'
    posts_url = f"{base.rstrip('/')}/wp-json/wp/v2/posts"
    response = requests.get(posts_url, headers=headers, params=params)
    if response.status_code >= 400:
        log.error(f"Failed to find post by slug '{slug}': {response.status_code} {response.text}")
        return None
    data = response.json()
    if isinstance(data, list) and data:
        return data[0]
    return None


def _confirm_update(existing, slug, force=False):
    if force:
        return True
    post_id = existing.get('id')
    status = existing.get('status') or ''
    date_val = existing.get('date') or ''
    prompt = f"Post with slug '{slug}' exists (id={post_id}, status={status}, date={date_val}). Update? [y/N]: "
    try:
        resp = input(prompt)
    except EOFError:
        return False
    return resp.strip().lower() in ('y', 'yes')


def schedule_post_wp_api(url, headers, content_md, meta, dry_run=False, preview=False, tz_name=DEFAULT_TIMEZONE, force=False, publish_date: Any | None = None):
    content_md = ensure_post_navigation_blocks(content_md)
    payload, schedule_dt = build_schedule_payload(url, headers, content_md, meta, tz_name, publish_date=publish_date)
    existing = find_post_by_slug(url, headers, payload.get('slug', ''))
    if existing:
        if not _confirm_update(existing, payload.get('slug', ''), force=force):
            log.info('Update cancelled by user.')
            existing_date = existing.get('date') or schedule_dt.isoformat()
            return {
                'action': 'skipped',
                'scheduled_for': existing_date,
                'payload': payload,
                'post': existing
            }
        if existing.get('status'):
            payload['status'] = existing.get('status')
        if existing.get('date') and not publish_date:
            payload['date'] = existing.get('date')
            schedule_dt = datetime.fromisoformat(existing.get('date'))
    exclude_ids = {existing.get('id')} if existing and existing.get('id') else None
    shifted = []
    if publish_date:
        publish_day = _coerce_publish_date(publish_date)
        shifted = shift_scheduled_posts(url, headers, publish_day, tz_name=tz_name, dry_run=dry_run, exclude_ids=exclude_ids)
        if shifted:
            log.info(f'+  Shifted {len(shifted)} scheduled post(s) by +1 day starting {publish_day.strftime("%m-%d-%Y")}')
    if dry_run:
        return {
            'scheduled_for': schedule_dt.isoformat(),
            'action': 'update' if existing else 'create',
            'payload': payload,
            'shifted_posts': format_shift_report(shifted)
        }
    posts_url = build_wp_api_posts_url(url)
    if existing:
        post_id = existing.get('id')
        update_url = posts_url.rstrip('/') + f'/{post_id}'
        response = requests.post(update_url, headers=headers, json=payload)
    else:
        response = requests.post(posts_url, headers=headers, json=payload)
    if response.status_code >= 400:
        raise ValueError(f"Post schedule failed: {response.status_code} {response.text}")
    post = response.json()
    if preview:
        return {
            'post': post,
            'scheduled_for': schedule_dt.isoformat(),
            'action': 'update' if existing else 'create',
            'payload': payload,
            'shifted_posts': format_shift_report(shifted)
        }
    return post

def fetch_wp_posts_page(api_posts_url, headers, page=1, per_page=WP_API_MAX_PER_PAGE, date_range=None, status=None):
    params = {
        'per_page': per_page,
        'page': page,
        'orderby': 'date',
        'order': 'desc',
    }
    if status:
        params['status'] = status
    else:
        params['status'] = 'publish'
    params = add_date_range_params(params, date_range)
    response = requests.get(api_posts_url, headers=headers, params=params)
    response.raise_for_status()
    posts = response.json()
    total_pages = int(response.headers.get('X-WP-TotalPages', 1))
    return posts, total_pages
def download_blog_posts_wp_api(url, number=None, outdir='posts', formats=None, indir=None, with_meta=False, date_range=None, status=None):
    '''
    Download blog posts using the WordPress REST API instead of scraping HTML pages.
    '''
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
    }

    # Set of already downloaded posts based on indir_list
    existing_txt_files = set()
    existing_docx_files = set()
    indir_list = indir if indir is not None else []
    for idir in (indir_list or []):
        if os.path.isdir(idir):
            for fname in os.listdir(idir):
                if fname.endswith(".txt"):
                    existing_txt_files.add(fname)
                elif fname.endswith(".docx"):
                    existing_docx_files.add(fname)

    if not os.path.exists(outdir):
        os.makedirs(outdir)

    # Track already-seen URLs
    seen_urls_path = os.path.join(outdir, 'seen_urls.txt')
    seen_urls = set()
    if os.path.exists(seen_urls_path):
        with open(seen_urls_path, 'r') as f:
            seen_urls = set(line.strip() for line in f if line.strip())

    api_posts_url = build_wp_api_posts_url(url)
    log.info(f'WordPress API endpoint: {api_posts_url}')

    formats = formats or ['txt']
    count = 0
    results = []
    stats = {
        'scanned': 0,
        'skipped_seen': 0,
        'skipped_duplicate': 0,
        'skipped_other': 0,
    }
    page = 1
    per_page = min(number, WP_API_MAX_PER_PAGE) if number else WP_API_MAX_PER_PAGE
    consecutive_duplicates = 0

    while True:
        try:
            posts, total_pages = fetch_wp_posts_page(
                api_posts_url,
                headers,
                page=page,
                per_page=per_page,
                date_range=date_range,
                status=status
            )
        except requests.RequestException:
            log.error(f"Failed to fetch WordPress API page: {api_posts_url} (page {page})")
            return results, stats

        if not posts:
            log.warning("No posts returned by WordPress API.")
            break

        for post in posts:
            stats['scanned'] += 1
            if number and count >= number:
                return results, stats

            post_url = post.get('link') or post.get('guid', {}).get('rendered')
            if post_url and post_url in seen_urls:
                log.info(f"Skipping already-seen URL: {post_url}")
                stats['skipped_seen'] += 1
                continue

            post_title_html = (post.get('title', {}) or {}).get('rendered', '')
            post_title = BeautifulSoup(post_title_html, 'html.parser').get_text().strip()
            if not post_title:
                log.warning("No title found in API post, skipping...")
                stats['skipped_other'] += 1
                continue

            publication_date = format_wp_api_date(post.get('date') or post.get('date_gmt'))
            content_html = (post.get('content', {}) or {}).get('rendered', '')
            content_text = extract_text_from_html(content_html, outdir)

            if not content_text:
                log.warning(f"No content found in {post_title}, skipping...")
                stats['skipped_other'] += 1
                continue

            sanitized_title = sanitize_filename(post_title)
            filename_with_date = f"{sanitized_title}-{publication_date}"
            text_filename = filename_with_date + '.txt'
            docx_filename = filename_with_date + '.docx'

            if is_duplicate_filename(text_filename, docx_filename, existing_txt_files, existing_docx_files):
                log.info(f"Skipping duplicate: {text_filename}")
                consecutive_duplicates += 1
                stats['skipped_duplicate'] += 1
                if consecutive_duplicates >= NUM_DUPLICATES:
                    log.info(f"Found {NUM_DUPLICATES} filename-based duplicates in a row. Stopping early.")
                    return results, stats
                continue
            else:
                consecutive_duplicates = 0

            text_path = None
            md_path = None
            docx_path = None
            if 'txt' in formats:
                text_path = save_text_file(post_title, publication_date, content_text, post_url or '', outdir, filename_with_date, 'txt')
            if 'md' in formats:
                md_path = save_text_file(post_title, publication_date, content_text, post_url or '', outdir, filename_with_date, 'md')
            if 'word' in formats:
                docx_path = save_word_file(post_title, publication_date, content_text, outdir, filename_with_date, True, log, post_url or '')
            meta_path = ''
            if with_meta:
                meta_path = save_meta_file(post, outdir, filename_with_date)
            results.append({
                'title': post_title,
                'date': publication_date,
                'url': post_url or '',
                'text_path': text_path or '',
                'md_path': md_path or '',
                'docx_path': docx_path or '',
                'meta_path': meta_path
            })

            if post_url:
                with open(seen_urls_path, 'a') as f:
                    f.write(post_url + '\n')

            count += 1

        if page >= total_pages:
            log.info('No more API pages to process.')
            break
        page += 1
    return results, stats

def extract_date_from_filename(filename):
    """
    Extracts and returns the date from a given filename in the format '[filename]-YYYY-MM-DD'.
    Returns None if the date format is invalid.
    """
    try:
        parts = filename.split('-')
        if len(parts) >= 3:
            date_str = '-'.join(parts[-3:])  # Extract the last three parts for YYYY-MM-DD
            date_str = date_str.split('.')[0]  # Remove file extension
            return datetime.strptime(date_str, "%Y-%m-%d").date()
    except ValueError:
        log.warning(f"Invalid date format in filename: {filename}")
    return None
            
def concatenate_text_files(indir_list, outfile, num_files, extension):
    '''
    Concatenate text files in the specified directories, sorted by date (newest first).
    Assumes that a function extract_date_from_filename(filename) is defined externally.
    '''
    files_with_dates = []
    for idir in indir_list:
        if not os.path.isdir(idir):
            continue
        for file in os.listdir(idir):
            if file.endswith(f'.{extension}'):
                file_date = extract_date_from_filename(file)
                if file_date:
                    files_with_dates.append((os.path.join(idir, file), file_date))
                else:
                    log.warning(f"Filename does not contain a valid date: {file}")

    # Sort files by date (newest first)
    files_with_dates.sort(key=lambda x: x[1], reverse=True)

    # Process files in sorted order
    with open(f'{outfile}.{extension}', 'w') as out_file:
        processed_files = 0
        for file_path, _ in files_with_dates:
            with open(file_path, 'r') as in_file:
                out_file.write(in_file.read() + '\n\n')
                log.info(f'Processed file: {os.path.basename(file_path)}')
                processed_files += 1

                # Break if num_files is reached
                if num_files is not None and processed_files >= num_files:
                    break

    log.info(f'Concatenated document saved as: {outfile}.{extension}')
    return processed_files


def concatenate_word_documents(indir_list, outfile, num_files=None):
    '''
    Concatenates Word documents from the specified directories into a single document.
    Properly preserves headings, paragraphs, and images.
    '''
    files_with_dates = []

    # Collect and sort files by date (newest first)
    for idir in indir_list:
        if not os.path.isdir(idir):
            continue
        for file in os.listdir(idir):
            if file.endswith('.docx'):
                file_date = extract_date_from_filename(file)
                if file_date:
                    files_with_dates.append((os.path.join(idir, file), file_date))
                else:
                    log.warning(f"Filename does not contain a valid date: {file}")

    # Sort files by date (newest first)
    files_with_dates.sort(key=lambda x: x[1], reverse=True)

    # Process files in sorted order
    outdoc = Document()
    processed_files = 0

    for file_path, _ in files_with_dates:
        src_doc = Document(file_path)

        # Add title as a heading (use filename without the date)
        title_text = os.path.basename(file_path).rsplit('-', 1)[0].replace('_', ' ')
        outdoc.add_heading(title_text, level=1)

        # Copy paragraphs and formatting from the original document
        i = 0
        while i < len(src_doc.paragraphs):
            paragraph = src_doc.paragraphs[i]
            # Preserve plain text of link instead of hyperlink
            if paragraph.text.strip() == "Original post link":
                if i + 1 < len(src_doc.paragraphs):
                    link_url = src_doc.paragraphs[i + 1].text.strip()
                    log.debug(f"Post link: {link_url}")
                    outdoc.add_paragraph(f"Original post link: {link_url}")
                i += 2
                continue
            new_paragraph = outdoc.add_paragraph(style=paragraph.style)
            for run in paragraph.runs:
                r = new_paragraph.add_run(run.text)
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline
            i += 1

        # Insert a page break between posts
        outdoc.add_page_break()

        log.info(f'Processed Word file: {os.path.basename(file_path)}')
        processed_files += 1

        if num_files is not None and processed_files >= num_files:
            break

    # Save the concatenated document
    outdoc.save(f'{outfile}.docx')
    log.info(f'Concatenated document saved as: {outfile}.docx')
    return processed_files

    
# ****************************************************************************************
# arguments
# ****************************************************************************************
def handle_args():
    load_dotenv()
    parser = argparse.ArgumentParser(description='Download blog posts as text documents')
    parser.add_argument(
        '--url', 
        default='johnmaconline.com',
        help='URL of the blog page (default: johnmaconline.com; scheme optional)')
    parser.add_argument(
        '--number', 
        type=int, 
        help='Number of blog posts to download')
    parser.add_argument(
        '--outdir', 
        default='posts', 
        help='Directory to save blog posts [default: posts]')    
    parser.add_argument(
        '--indir',
        nargs='+',
        help='One or more input dirs for file concat (space-separated)')
    parser.add_argument(
        '--outfile',
        nargs='?',
        const='',
        help='Output file base name (defaults to out.<format> if omitted)')    
    parser.add_argument(
        '--outfile-format',
        choices=['csv', 'json'],
        default='csv',
        help='Output format for results file [default: csv]')
    parser.add_argument(
        '--format',
        nargs='+',
        choices=['txt', 'md', 'word'],
        default=['txt'],
        help='Output formats for downloaded posts [default: txt]')
    parser.add_argument(
        '--date',
        help='Date range filter: today, week, month, year, all, or MM-DD-YYYY:MM-DD-YYYY (overrides --number)')
    parser.add_argument(
        '--post-state',
        choices=['published', 'scheduled', 'draft'],
        default='published',
        help='Post state to fetch [default: published]')
    parser.add_argument(
        '--get-posts',
        action='store_true',
        help='Download posts via the REST API.')
    parser.add_argument(
        '--schedule-post',
        action='store_true',
        help='Schedule a post via the REST API using markdown content and metadata JSON.')
    parser.add_argument(
        '--backfill-post-navigation',
        action='store_true',
        help='Append Gutenberg previous/next navigation blocks to existing posts that do not already have them.')
    parser.add_argument(
        '--post-id',
        type=int,
        help='Specific WordPress post ID to target for --backfill-post-navigation.')
    parser.add_argument(
        '--content-md',
        help='Path to markdown content file (use - for stdin) for --schedule-post')
    parser.add_argument(
        '--meta-json',
        help='Path to metadata JSON file for --schedule-post')
    parser.add_argument(
        '--publish-date',
        help='Publish date for --schedule-post (MM-DD-YYYY, scheduled at 8:44am Eastern)')
    parser.add_argument(
        '--dry-run',
        action='store_true',
        help='Show planned changes without making updates (used with --schedule-post).')
    parser.add_argument(
        '--force',
        action='store_true',
        help='Skip confirmation prompts (used with --schedule-post).')
    parser.add_argument(
        '--preview',
        action='store_true',
        help='Print the final scheduled post payload (used with --schedule-post).')
    parser.add_argument(
        '--export-site',
        help='Export site data via REST API: all, all-no-media, or comma-separated list (posts,pages,media,comments,users,categories,tags,taxonomies,types,statuses,settings,menus,plugins)')
    parser.add_argument(
        '--incremental',
        action='store_true',
        help='For --export-site, append only new items to JSONL outputs.')
    parser.add_argument(
        '--get-plugins',
        action='store_true',
        help='Fetch and print WordPress plugins via the REST API (requires credentials).')
    parser.add_argument(
        '--list-posts',
        action='store_true',
        help='Fetch and print posts via the REST API.')
    parser.add_argument(
        '--get-pages',
        action='store_true',
        help='Fetch and print pages via the REST API.')
    parser.add_argument(
        '--get-categories',
        action='store_true',
        help='Fetch and print categories via the REST API.')
    parser.add_argument(
        '--get-tags',
        action='store_true',
        help='Fetch and print tags via the REST API.')
    parser.add_argument(
        '--get-users',
        action='store_true',
        help='Fetch and print users via the REST API.')
    parser.add_argument(
        '--get-user-me',
        action='store_true',
        help='Fetch and print current user details (requires credentials).')
    parser.add_argument(
        '--get-media',
        action='store_true',
        help='Fetch and print media items via the REST API.')
    parser.add_argument(
        '--get-comments',
        action='store_true',
        help='Fetch and print comments via the REST API.')
    parser.add_argument(
        '--get-types',
        action='store_true',
        help='Fetch and print content types via the REST API.')
    parser.add_argument(
        '--get-statuses',
        action='store_true',
        help='Fetch and print post statuses via the REST API.')
    parser.add_argument(
        '--get-taxonomies',
        action='store_true',
        help='Fetch and print taxonomies via the REST API.')
    parser.add_argument(
        '--get-settings',
        action='store_true',
        help='Fetch and print site settings (requires credentials).')
    parser.add_argument(
        '--get-themes',
        action='store_true',
        help='Fetch and print themes (requires credentials).')
    parser.add_argument(
        '--wp-username',
        default=os.getenv('WP_USERNAME'),
        help='WordPress username (or set WP_USERNAME env var)')
    parser.add_argument(
        '--wp-app-password',
        default=os.getenv('WP_APP_PASSWORD'),
        help='WordPress application password (or set WP_APP_PASSWORD env var)')
    parser.add_argument(
        '--with-meta',
        action='store_true',
        help='Save per-post metadata JSON alongside downloaded posts.')
    parser.add_argument(
        '--concat',
        action='store_true',
        help='Concat files.')
    parser.add_argument(
        '-v',
        '--verbose',
        action='store_true',
        help='Enable verbose output to stdout.')
    parser.add_argument(
        '-q',
        '--quiet',
        action='store_true',
        help='Minimal stdout.')
    args = parser.parse_args()

    # Configure stdout logging based on arguments
    ch = logging.StreamHandler(sys.stdout)
    if args.verbose:
        ch.setLevel(logging.DEBUG)
    elif args.quiet:
        ch.setLevel(logging.ERROR)
    else:
        ch.setLevel(logging.INFO)
    ch.setFormatter(formatter)
    log.addHandler(ch)
    try:
        export_items, export_download_media = parse_export_site_arg(args.export_site) if args.export_site else (set(), False)
    except ValueError as exc:
        log.error(str(exc))
        sys.exit(1)
    args.export_site_items = export_items
    args.export_download_media = export_download_media
    
    # check requirements
    if args.get_posts:
        if args.url:
            log.debug(f'Download requirements met')
        else:
            log.error(f'XXX Must supply a URL')
            sys.exit(1)
    if args.date and args.number:
        log.info('++ --date provided; ignoring --number')
    if args.post_state != 'published' and not (args.wp_username and args.wp_app_password):
        log.error('XXX Non-published post states require WP credentials')
        sys.exit(1)
    if args.export_site and not args.url:
        log.error('XXX Must supply a URL for --export-site')
        sys.exit(1)
    if args.schedule_post:
        if not args.url:
            log.error('XXX Must supply a URL for --schedule-post')
            sys.exit(1)
        if not args.content_md or not args.meta_json:
            log.error('XXX Must supply --content-md and --meta-json for --schedule-post')
            sys.exit(1)
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --schedule-post')
            sys.exit(1)
        if args.publish_date:
            try:
                parse_publish_date(args.publish_date)
            except ValueError as exc:
                log.error(f'XXX {exc}')
                sys.exit(1)
    if args.backfill_post_navigation:
        if not args.url:
            log.error('XXX Must supply a URL for --backfill-post-navigation')
            sys.exit(1)
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --backfill-post-navigation')
            sys.exit(1)
    if args.post_id and not args.backfill_post_navigation:
        log.error('XXX --post-id requires --backfill-post-navigation')
        sys.exit(1)
    if (args.get_plugins or args.list_posts or args.get_pages or args.get_categories or args.get_tags or
            args.get_users or args.get_user_me or args.get_media or args.get_comments or args.get_types or
            args.get_statuses or args.get_taxonomies or args.get_settings or args.get_themes) and not args.url:
        log.error('XXX Must supply a URL for --get-* operations')
        sys.exit(1)

    log.info('++++++++++++++++++++++++++++++++++++++++++++++')
    log.info(f'+  {os.path.basename(sys.argv[0])}')
    log.info(f'+  Python Version: {sys.version.split()[0]}')
    log.info(f'+  Today is: {date.today()}')
    if args.export_site:
        log.info('+  Exporting site via WP API')
        log.info(f'+  Target URL: {args.url}')
        log.info(f'+  Export directory: {EXPORT_DEFAULT_DIR}')
        log.info(f'+  Export items: {", ".join(sorted(args.export_site_items))}')
        log.info(f'+  Incremental: {args.incremental}')
        if 'media' in args.export_site_items:
            if args.export_download_media:
                log.info('+  Media binaries: enabled')
            else:
                log.info('+  Media binaries: disabled (metadata only)')
        if args.export_site_items.intersection({'users', 'settings'}) and not (args.wp_username and args.wp_app_password):
            log.info('+  Credentials not provided; some export items may fail')
    if args.schedule_post:
        log.info('+  Scheduling post via WP API')
        log.info(f'+  Target URL: {args.url}')
        log.info(f'+  Content markdown: {args.content_md}')
        log.info(f'+  Metadata JSON: {args.meta_json}')
        log.info(f'+  Timezone: {DEFAULT_TIMEZONE}')
        if args.publish_date:
            log.info(f'+  Publish date: {args.publish_date} (8:44am Eastern)')
        if args.dry_run:
            log.info('+  Dry run enabled (no post will be created)')
        if args.preview:
            log.info('+  Preview enabled (payload will be printed)')
        if args.force:
            log.info('+  Force enabled (skip update prompts)')
    if args.backfill_post_navigation:
        log.info('+  Backfilling post navigation blocks via WP API')
        log.info(f'+  Target URL: {args.url}')
        if args.post_id is not None:
            log.info(f'+  Post ID: {args.post_id}')
        else:
            log.info(f'+  Post state: {args.post_state}')
        if args.number is not None:
            log.info(f'+  Limit: {args.number}')
        if args.dry_run:
            log.info('+  Dry run enabled (no posts will be updated)')
    if args.get_posts:
        log.info('+  Fetching posts via WP API')
        log.info(f'+  Target URL: {args.url}')
        log.info(f'+  Post state: {args.post_state}')
        if args.date:
            log.info(f'+  Date range: {args.date}')
        else:
            log.info(f'+  Number of posts to download: {"All" if args.number is None else args.number}')
        log.info(f'+  Output directory: {args.outdir}')
        log.info(f'+  Formats: {args.format}')
        if args.with_meta:
            log.info('+  Downloading metadata')
        log.info('+  Downloading post content')
    if args.get_plugins:
        log.info('+  Fetching plugin list via WP API')
    if args.list_posts:
        log.info('+  Fetching posts via WP API')
        if args.date and not args.get_posts:
            log.info(f'+  Date range: {args.date}')
        log.info(f'+  Post state: {args.post_state}')
    if args.get_pages:
        log.info('+  Fetching pages via WP API')
    if args.get_categories:
        log.info('+  Fetching categories via WP API')
    if args.get_tags:
        log.info('+  Fetching tags via WP API')
    if args.get_users:
        log.info('+  Fetching users via WP API')
    if args.get_user_me:
        log.info('+  Fetching current user via WP API')
    if args.get_media:
        log.info('+  Fetching media via WP API')
    if args.get_comments:
        log.info('+  Fetching comments via WP API')
    if args.get_types:
        log.info('+  Fetching types via WP API')
    if args.get_statuses:
        log.info('+  Fetching statuses via WP API')
    if args.get_taxonomies:
        log.info('+  Fetching taxonomies via WP API')
    if args.get_settings:
        log.info('+  Fetching settings via WP API')
    if args.get_themes:
        log.info('+  Fetching themes via WP API')
    results_ops = [
        args.get_posts, args.list_posts, args.get_plugins, args.get_pages, args.get_categories, args.get_tags,
        args.get_users, args.get_user_me, args.get_media, args.get_comments, args.get_types, args.get_statuses,
        args.get_taxonomies, args.get_settings, args.get_themes, args.schedule_post, args.backfill_post_navigation
    ]
    if any(results_ops):
        if not args.get_posts and args.url:
            log.info(f'+  Target URL: {args.url}')
        actions = []
        if args.get_posts:
            actions.append('get-posts (download)')
        if args.schedule_post:
            actions.append('schedule-post')
        if args.backfill_post_navigation:
            actions.append('backfill-post-navigation')
        if args.list_posts:
            actions.append('list-posts')
        if args.get_plugins:
            actions.append('get-plugins')
        if args.get_pages:
            actions.append('get-pages')
        if args.get_categories:
            actions.append('get-categories')
        if args.get_tags:
            actions.append('get-tags')
        if args.get_users:
            actions.append('get-users')
        if args.get_user_me:
            actions.append('get-user-me')
        if args.get_media:
            actions.append('get-media')
        if args.get_comments:
            actions.append('get-comments')
        if args.get_types:
            actions.append('get-types')
        if args.get_statuses:
            actions.append('get-statuses')
        if args.get_taxonomies:
            actions.append('get-taxonomies')
        if args.get_settings:
            actions.append('get-settings')
        if args.get_themes:
            actions.append('get-themes')
        if len(actions) > 1:
            log.info(f'+  Actions: {", ".join(actions)}')

        results_outfile = args.outfile or f"out.{args.outfile_format}"
        log.info(f'+  Results summary file: {results_outfile}')
        log.info(f'+  Results format: {args.outfile_format}')
    if args.concat:
        log.info(f'+  Concatenating files')
        log.info(f'+  Output file: {args.outfile}')
        concat_outputs = []
        if 'txt' in args.format:
            concat_outputs.append(f"{args.outfile}.txt")
        if 'md' in args.format:
            concat_outputs.append(f"{args.outfile}.md")
        if 'word' in args.format:
            concat_outputs.append(f"{args.outfile}.docx")
        if concat_outputs:
            log.info(f'+  Concat outputs: {", ".join(concat_outputs)}')
        if not args.get_posts:
            if args.indir:
                log.info(f'+  Concatenating from: {args.indir}')
            else:
                log.error(f'XXX Must include --indir')
                sys.exit(1)
    log.info('++++++++++++++++++++++++++++++++++++++++++++++')        

    return args

# ****************************************************************************************
# Main
# ****************************************************************************************
def main():
    args = handle_args()
    if args.dry_run:
        atexit.register(lambda: log.info(f'NO changes were made to the WP site: {args.url}'))
    # Build initial list of directories to check for existing posts
    indir_list = args.indir or []
    results_outfile = args.outfile or f"out.{args.outfile_format}"
    results_rows = []
    op_counts = {}
    summary = {
        'requested_posts': 0,
        'scanned_posts': 0,
        'download_posts': 0,
        'download_txt': 0,
        'download_md': 0,
        'download_docx': 0,
        'download_meta': 0,
        'skipped_seen': 0,
        'skipped_duplicate': 0,
        'skipped_other': 0,
        'concat_txt': 0,
        'concat_md': 0,
        'concat_docx': 0,
    }
    results_file_path = None
    export_summary = None
    date_range = (None, None)
    if args.date:
        try:
            date_range = parse_date_range(args.date)
        except ValueError as exc:
            log.error(str(exc))
            sys.exit(1)
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
    }
    if args.wp_username and args.wp_app_password:
        headers.update(build_auth_header(args.wp_username, args.wp_app_password))
    if args.export_site:
        export_summary = export_site(
            args.url,
            headers,
            args.export_site_items,
            EXPORT_DEFAULT_DIR,
            args.incremental,
            args.export_download_media
        )
    if args.schedule_post:
        try:
            meta = load_meta_json(args.meta_json)
            content_md = strip_leading_h1(read_markdown_content(args.content_md))
            base_dir = os.path.dirname(args.content_md) if args.content_md != '-' else os.getcwd()
            uploads = []
            if args.dry_run:
                log.info('*** this is a dry-run ***')
                local_images = find_local_markdown_images(content_md, base_dir)
                missing = [path for path in local_images if not os.path.exists(path)]
                if missing:
                    raise ValueError(f"Missing local images: {', '.join(missing)}")
                if local_images:
                    log.info(f'+  Dry run: would upload {len(local_images)} image(s)')
                result = schedule_post_wp_api(
                    args.url,
                    headers,
                    content_md,
                    meta,
                    dry_run=True,
                    preview=args.preview,
                    tz_name=DEFAULT_TIMEZONE,
                    force=args.force,
                    publish_date=args.publish_date
                )
                shifted_posts = result.get('shifted_posts') or []
                if args.publish_date:
                    shifted_posts = list(shifted_posts)
                    shifted_posts.append({
                        'operation': 'shift-scheduled',
                        'id': '',
                        'title': meta.get('title') or meta.get('post_title') or '',
                        'from': '(new)',
                        'to': result.get('scheduled_for') or ''
                    })
                if shifted_posts:
                    log.info('+  Shift report (dry run)')
                    print(render_ascii_table(shifted_posts))
                action = result.get('action')
                if action == 'skipped':
                    log.info('Update skipped.')
                    post = result.get('post', {})
                    scheduled_for = result.get('scheduled_for')
                    post_id = post.get('id') or ''
                    link = post.get('link') or ''
                else:
                    if args.preview:
                        print(json.dumps(result.get('payload', {}), indent=2))
                    scheduled_for = result.get('scheduled_for')
                    post_id = ''
                    link = ''
            else:
                content_md, uploads = upload_media_and_replace(content_md, base_dir, args.url, headers)
                if uploads:
                    log.info(f'+  Uploaded {len(uploads)} image(s)')
                result = schedule_post_wp_api(
                    args.url,
                    headers,
                    content_md,
                    meta,
                    dry_run=False,
                    preview=args.preview,
                    tz_name=DEFAULT_TIMEZONE,
                    force=args.force,
                    publish_date=args.publish_date
                )
                if args.preview:
                    shifted_posts = result.get('shifted_posts') or []
                    if shifted_posts:
                        log.info('+  Shift report')
                        print(render_ascii_table(shifted_posts))
                action = result.get('action')
                if action == 'skipped':
                    log.info('Update skipped.')
                    post = result.get('post', {})
                else:
                    if args.preview:
                        print(json.dumps(result.get('payload', {}), indent=2))
                        post = result.get('post', {})
                    else:
                        post = result
                scheduled_for = post.get('date') or result.get('scheduled_for')
                post_id = post.get('id')
                link = post.get('link') or ''
            title = meta.get('title') or meta.get('post_title') or ''
            row = {
                'operation': 'schedule-post',
                'title': title,
                'scheduled_for': scheduled_for or '',
                'status': 'future',
                'post_id': post_id or '',
                'link': link,
                'media_uploaded': len(uploads),
                'dry_run': str(bool(args.dry_run)),
                'action': action or ''
            }
            results_rows.append(row)
            op_counts['schedule-post'] = 1
            print(render_ascii_table([row]))
        except ValueError as exc:
            log.error(f'XXX {exc}')
            sys.exit(1)
    if args.backfill_post_navigation:
        try:
            status_map = {'published': 'publish', 'scheduled': 'future', 'draft': 'draft'}
            rows, stats = backfill_post_navigation(
                args.url,
                headers,
                dry_run=args.dry_run,
                status=status_map.get(args.post_state, 'publish'),
                limit=args.number,
                post_id=args.post_id,
            )
            op_counts['backfill-post-navigation'] = len(rows)
            results_rows.extend(rows)
            if rows:
                print(render_ascii_table(rows))
            log.info(
                '+  Backfill summary: '
                f'processed={stats["processed"]}, '
                f'already_has_navigation={stats["already_has_navigation"]}, '
                f'needs_update={stats["needs_update"]}, '
                f'skipped={stats["skipped"]}'
            )
        except ValueError as exc:
            log.error(f'XXX {exc}')
            sys.exit(1)
    if args.get_plugins:
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --get-plugins')
            sys.exit(1)
        plugins = fetch_wp_plugins(args.url, headers)
        if plugins is None:
            sys.exit(1)
        rows = []
        for plugin in plugins:
            rows.append({
                'operation': 'get-plugins',
                'name': (plugin.get('name') or ''),
                'status': plugin.get('status') or '',
                'version': plugin.get('version') or '',
                'plugin': plugin.get('plugin') or ''
            })
        op_counts['get-plugins'] = len(rows)
        results_rows.extend(rows)
        print(render_ascii_table(rows))
    if args.list_posts:
        list_params = {}
        if args.date:
            list_params = add_date_range_params(list_params, date_range)
        if args.post_state:
            status_map = {'published': 'publish', 'scheduled': 'future', 'draft': 'draft'}
            list_params['status'] = status_map.get(args.post_state, 'publish')
        data = fetch_wp_endpoint(args.url, 'posts', headers, params=list_params)
        rows = normalize_wp_rows('list-posts', data)
        op_counts['list-posts'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_pages:
        data = fetch_wp_endpoint(args.url, 'pages', headers)
        rows = normalize_wp_rows('get-pages', data)
        op_counts['get-pages'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_categories:
        data = fetch_wp_endpoint(args.url, 'categories', headers)
        rows = normalize_wp_rows('get-categories', data)
        op_counts['get-categories'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_tags:
        data = fetch_wp_endpoint(args.url, 'tags', headers)
        rows = normalize_wp_rows('get-tags', data)
        op_counts['get-tags'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_users:
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --get-users')
            sys.exit(1)
        data = fetch_wp_endpoint(args.url, 'users', headers, params={'context': 'edit'})
        rows = normalize_wp_rows('get-users', data)
        op_counts['get-users'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_user_me:
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --get-user-me')
            sys.exit(1)
        data = fetch_wp_endpoint(args.url, 'users/me', headers, params={'context': 'edit'})
        rows = normalize_wp_rows('get-user-me', data)
        op_counts['get-user-me'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_media:
        data = fetch_wp_endpoint(args.url, 'media', headers)
        rows = normalize_wp_rows('get-media', data)
        op_counts['get-media'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_comments:
        data = fetch_wp_endpoint(args.url, 'comments', headers)
        rows = normalize_wp_rows('get-comments', data)
        op_counts['get-comments'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_types:
        data = fetch_wp_endpoint(args.url, 'types', headers)
        rows = normalize_wp_rows('get-types', data)
        op_counts['get-types'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_statuses:
        data = fetch_wp_endpoint(args.url, 'statuses', headers)
        rows = normalize_wp_rows('get-statuses', data)
        op_counts['get-statuses'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_taxonomies:
        data = fetch_wp_endpoint(args.url, 'taxonomies', headers)
        rows = normalize_wp_rows('get-taxonomies', data)
        op_counts['get-taxonomies'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_settings:
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --get-settings')
            sys.exit(1)
        data = fetch_wp_endpoint(args.url, 'settings', headers)
        rows = normalize_wp_rows('get-settings', data)
        op_counts['get-settings'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    if args.get_themes:
        if not args.wp_username or not args.wp_app_password:
            log.error('XXX Missing WP credentials for --get-themes')
            sys.exit(1)
        data = fetch_wp_endpoint(args.url, 'themes', headers)
        rows = normalize_wp_rows('get-themes', data)
        op_counts['get-themes'] = len(rows)
        results_rows.extend(rows)
        if rows:
            print(render_ascii_table(rows))
    # Download, using existing indices from all provided indir_list
    if args.get_posts:
        summary['requested_posts'] = args.number or 0
        status_map = {'published': 'publish', 'scheduled': 'future', 'draft': 'draft'}
        status = status_map.get(args.post_state, 'publish')
        download_results, download_stats = download_blog_posts_wp_api(
            args.url,
            None if args.date else args.number,
            args.outdir,
            args.format,
            indir_list,
            args.with_meta,
            date_range,
            status
        )
        for row in (download_results or []):
            row['operation'] = 'get-posts'
            results_rows.append(row)
        if download_results:
            summary['download_posts'] = len(download_results)
            summary['download_txt'] = sum(1 for r in download_results if r.get('text_path'))
            summary['download_md'] = sum(1 for r in download_results if r.get('md_path'))
            summary['download_docx'] = sum(1 for r in download_results if r.get('docx_path'))
            summary['download_meta'] = sum(1 for r in download_results if r.get('meta_path'))
            op_counts['get-posts'] = len(download_results)
            print(render_ascii_table(download_results))
        if download_stats:
            summary['scanned_posts'] = download_stats.get('scanned', 0)
            summary['skipped_seen'] = download_stats.get('skipped_seen', 0)
            summary['skipped_duplicate'] = download_stats.get('skipped_duplicate', 0)
            summary['skipped_other'] = download_stats.get('skipped_other', 0)
        # After downloading into outdir, include it for concat
        indir_list = indir_list + [args.outdir]
    # Concatenate from all directories (original and newly downloaded)
    if args.concat:
        if 'txt' in args.format:
            summary['concat_txt'] = concatenate_text_files(indir_list, args.outfile, args.number, 'txt')
        if 'md' in args.format:
            summary['concat_md'] = concatenate_text_files(indir_list, args.outfile, args.number, 'md')
        if 'word' in args.format:
            summary['concat_docx'] = concatenate_word_documents(indir_list, args.outfile, args.number)
    if results_rows:
        out_path = write_results(results_outfile, args.outfile_format, results_rows)
        if out_path:
            log.info(f'Wrote results file: {out_path}')
            results_file_path = out_path
    log.info('+++++++++++++++++++++++++ Summary +++++++++++++++++++++++++')
    if op_counts:
        for op_name, count in sorted(op_counts.items()):
            log.info(f'+  {op_name}: {count}')
    if results_rows:
        log.info(f'+  Rows written to file: {len(results_rows)}')
        if results_file_path:
            log.info(f'+  Results file path: {results_file_path}')
    if summary['download_posts']:
        if args.date:
            log.info(f'+  Date range: {args.date}')
        elif summary['requested_posts']:
            log.info(f'+  Requested downloads: {summary["requested_posts"]}')
        if summary['scanned_posts']:
            log.info(f'+  Scanned posts: {summary["scanned_posts"]}')
        log.info(f'+  Downloaded posts: {summary["download_posts"]}')
        log.info(f'+  Files written: txt={summary["download_txt"]}, md={summary["download_md"]}, word={summary["download_docx"]}, meta={summary["download_meta"]}')
        if args.with_meta:
            log.info(f'+  Metadata files path: {args.outdir}/*.meta.json')
        skipped_total = summary['skipped_seen'] + summary['skipped_duplicate'] + summary['skipped_other']
        log.info(f'+  Skipped posts: {skipped_total} (seen={summary["skipped_seen"]}, duplicate={summary["skipped_duplicate"]}, other={summary["skipped_other"]})')
    if args.concat:
        concat_parts = []
        if 'txt' in args.format:
            concat_parts.append(f"txt={summary['concat_txt']}")
        if 'md' in args.format:
            concat_parts.append(f"md={summary['concat_md']}")
        if 'word' in args.format:
            concat_parts.append(f"word={summary['concat_docx']}")
        if concat_parts:
            log.info(f'+  Concat files written: {", ".join(concat_parts)}')
    if export_summary:
        log.info(f'+  Export output: {EXPORT_DEFAULT_DIR}')
        resources = export_summary.get('resources', {})
        for resource_name in sorted(resources.keys()):
            info = resources[resource_name]
            if resource_name == 'media_files':
                log.info(
                    f'+  Media files: downloaded={info.get("downloaded", 0)}, '
                    f'skipped_existing={info.get("skipped_existing", 0)}, '
                    f'skipped_missing_url={info.get("skipped_missing_url", 0)}, '
                    f'failed={info.get("failed", 0)}'
                )
                log.info(f'+  Media files path: {info.get("path", "")}')
                continue
            if 'written' in info:
                log.info(
                    f'+  Exported {resource_name}: total={info.get("count", 0)}, '
                    f'written={info.get("written", 0)}, '
                    f'skipped_existing={info.get("skipped_existing", 0)}'
                )
            else:
                log.info(f'+  Exported {resource_name}: count={info.get("count", 0)}')
            if info.get('path'):
                log.info(f'+  {resource_name} file: {info.get("path")}')
        if export_summary.get('manifest'):
            log.info(f'+  Export manifest: {export_summary["manifest"]}')
        warnings = export_summary.get('warnings') or []
        if warnings:
            log.info('+  Export warnings:')
            for warning in warnings:
                log.info(f'+    {warning}')
    log.info('+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++')

if __name__ == '__main__':
    main()
