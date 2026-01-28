##########################################################################################
#
# Script name: wp_utilities.py
#
# Description: Downloads a specified number of blog posts as text documents from a given URL.
#              Each blog post is saved as a text file with the blog title as the filename
#              in the specified output directory.
#
# Author: [Your Name]
#
##########################################################################################

import requests
from bs4 import BeautifulSoup, NavigableString
import argparse
import logging
import sys
import os
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from datetime import datetime
from urllib.parse import urlparse, unquote

# ****************************************************************************************
# Global data and configuration
# ****************************************************************************************

# Logging config
log = logging.getLogger(os.path.basename(sys.argv[0]))
log.setLevel(logging.DEBUG)

# File handler for logging
fh = logging.FileHandler('wp_utilities.log', mode='w')
fh.setLevel(logging.DEBUG)
formatter = logging.Formatter(
    '%(asctime)-15s [%(funcName)25s:%(lineno)-5s] %(levelname)-8s %(message)s')
fh.setFormatter(formatter)
log.addHandler(fh)

# Number of consecutive filename-based duplicates before stopping early
NUM_DUPLICATES = 5
# WordPress API max per_page is typically 100
WP_API_MAX_PER_PAGE = 100

# ****************************************************************************************
# Exceptions
# ****************************************************************************************

class Error(Exception):
    """Base class for exceptions in this module."""
    pass

class RequestError(Error):
    def __init__(self, url):
        self.message = f"Failed to fetch URL: {url}"
        super().__init__(self.message)

# ****************************************************************************************
# Functions
# ****************************************************************************************
def valid_filename(url):
    """
    Create a valid filename from a URL by keeping only alphanumeric characters and replacing others,
    while preserving the original file extension.
    """
    parsed_url = urlparse(url)
    filename, file_extension = os.path.splitext(os.path.basename(parsed_url.path))
    # Keep alphanumeric characters and replace others with underscore
    filename = re.sub(r'\W+', '_', unquote(filename))
    # Concatenate filename and file extension
    return filename + file_extension


def download_image(image_url, outdir):
    '''
    Download an image and save it to the specified subdirectory ('images') within outdir.
    Returns the local filename of the downloaded image.
    '''
    images_dir = os.path.join(outdir, 'images')
    if not os.path.exists(images_dir):
        os.makedirs(images_dir)

    # Extract base URL (without query parameters) and create a valid filename
    base_url = image_url.split('?')[0]
    local_filename = os.path.join(images_dir, valid_filename(base_url))

    try:
        response = requests.get(image_url, stream=True)
        if response.status_code == 200:
            with open(local_filename, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            # Verify the downloaded file
            if os.path.getsize(local_filename) > 0:
                return local_filename
            else:
                log.error(f"Downloaded image is empty: {image_url}")
        else:
            log.error(f"Failed to download image: {image_url}, Status code: {response.status_code}")
    except Exception as e:
        log.error(f"Error downloading image {image_url}: {e}")
        
    log.debug(f"Image downloaded and saved to: {local_filename}")


    return None


def extract_text_from_html(html_content, outdir):
    '''
    Extract text from HTML content (WordPress API) including headings, lists, code blocks, and images.
    '''
    soup = BeautifulSoup(html_content or "", 'html.parser')
    container = soup.body if soup.body else soup
    content_text = ''

    def append_code_block(code_text):
        nonlocal content_text
        if not code_text:
            return
        code_text = code_text.strip('\n')
        if not code_text.strip():
            return
        content_text += f"```\n{code_text}\n```\n\n"

    def append_image(img_tag):
        nonlocal content_text
        img_url = img_tag.get('src') or img_tag.get('data-src')
        if img_url:
            local_image = download_image(img_url, outdir)
            if local_image:
                content_text += f'[image: {local_image}]\n'
            else:
                content_text += f'[image: {img_url} (download failed)]\n'
            content_text += '\n'

    def append_list(list_tag, ordered=False):
        nonlocal content_text
        index = 1
        for li in list_tag.find_all('li', recursive=False):
            item_text = li.get_text(strip=True)
            if ordered:
                content_text += f"{index}. {item_text}\n"
                index += 1
            else:
                content_text += f"* {item_text}\n"
        content_text += '\n'

    def walk(node):
        nonlocal content_text
        if isinstance(node, NavigableString):
            return
        if not hasattr(node, 'name'):
            return
        if node.name in {'script', 'style', 'noscript'}:
            return

        if node.name == 'h1':
            content_text += f"# {node.get_text(strip=True)}\n\n"
            return
        if node.name == 'h2':
            content_text += f"## {node.get_text(strip=True)}\n\n"
            return
        if node.name == 'p':
            text = node.get_text(strip=True)
            if text:
                content_text += text + '\n\n'
            return
        if node.name in {'pre', 'code'}:
            if node.name == 'code' and node.parent and node.parent.name == 'pre':
                return
            append_code_block(node.get_text())
            return
        if node.name == 'div':
            classes = node.get('class', [])
            if 'wp-block-code' in classes or 'wp-block-preformatted' in classes:
                pre = node.find('pre')
                code = node.find('code') if pre is None else None
                if pre:
                    append_code_block(pre.get_text())
                elif code:
                    append_code_block(code.get_text())
                return
        if node.name == 'ul':
            append_list(node, ordered=False)
            return
        if node.name == 'ol':
            append_list(node, ordered=True)
            return
        if node.name == 'figure':
            classes = node.get('class', [])
            if 'wp-block-image' in classes or node.find('img'):
                img_tag = node.find('img')
                if img_tag:
                    append_image(img_tag)
                return
        if node.name == 'img':
            append_image(node)
            return

        classes = node.get('class', [])
        if classes and 'social-sharing' in classes:
            return

        for child in node.children:
            walk(child)

    for child in container.children:
        walk(child)

    return content_text


def create_word_document(post_title, publish_date, content_text, outdir, post_url):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    def ensure_code_style(doc):
        try:
            return doc.styles['CodeBlock']
        except KeyError:
            style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Consolas'
            style.font.size = Pt(10)
            style.paragraph_format.left_indent = Inches(0.25)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            return style

    def add_code_paragraph(doc, lines):
        if not lines:
            return
        code_style = ensure_code_style(doc)
        paragraph = doc.add_paragraph(style=code_style)
        for i, line in enumerate(lines):
            if i:
                paragraph.add_run().add_break()
            paragraph.add_run(line)

    doc = Document()
    doc.add_heading(post_title, level=0)
    doc.add_paragraph(publish_date)

    in_code_block = False
    code_lines = []
    for line in content_text.split('\n'):
        if line.strip() == '```':
            if in_code_block:
                add_code_paragraph(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(line)
            continue
        if '[image:' in line:
            image_path = line[line.find('[image:') + 7:].rstrip(']').strip()
            full_image_path = os.path.join(outdir, image_path)
            if os.path.exists(full_image_path) and os.path.getsize(full_image_path) > 0:
                try:
                    doc.add_picture(full_image_path, width=Pt(350))
                except Exception as e:
                    log.error(f"Failed to embed image {full_image_path} in document: {e}")
                    doc.add_paragraph(f"[Error embedding image: {image_path}]")
        elif line.strip():
            doc.add_paragraph(line)

    if in_code_block and code_lines:
        add_code_paragraph(doc, code_lines)

    # Insert post link as a marker paragraph, followed by the URL (moved to end)
    doc.add_paragraph("Original post link")
    doc.add_paragraph(post_url)

    return doc

def sanitize_filename(title):
    '''
    Sanitizes the title to be safe for use as a filename.
    '''
    # Replace all spaces with underscores
    log.debug(f'+++ Sanitizing title: {title}')
    t1 = title.replace(' ', '_')
    t2 = t1.replace("’", "")
    san_title = re.sub(r'[^\w_]', '-', t2)
    log.debug(f'san_title: {san_title}')
    return san_title


# Helper: check for duplicate filenames
def is_duplicate_filename(text_filename, docx_filename, existing_txt_files, existing_docx_files):
    """
    Returns True if either the text or docx filename already exists in the respective sets.
    """
    return text_filename in existing_txt_files or docx_filename in existing_docx_files

# Helper: save text file
def save_text_file(post_title, publication_date, content_text, post_url, outdir, filename_with_date):
    """
    Save the blog post as a Markdown text file.
    """
    md_text = f"# {post_title}\nDate: {publication_date}\n\n{content_text}"
    text_filepath = os.path.join(outdir, filename_with_date + '.txt')
    with open(text_filepath, 'w') as file:
        file.write(md_text + f"\nOriginal post: {post_url}\n")
    log.info(f'Article saved as: {text_filepath}')
    return text_filepath

# Helper: save Word file
def save_word_file(post_title, publication_date, content_text, outdir, filename_with_date, word, log, post_url):
    """
    Save the blog post as a Word document, if requested.
    """
    if word:
        doc_title = post_title.replace('_', ' ').title()
        doc = create_word_document(doc_title, publication_date, content_text, outdir, post_url)
        word_filepath = os.path.join(outdir, filename_with_date + '.docx')
        os.makedirs(os.path.dirname(word_filepath), exist_ok=True)
        doc.save(word_filepath)
        log.info(f'Article saved as: {word_filepath}')
        return word_filepath
    return None


# Helper function to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.clear_content()
    paragraph._p.append(hyperlink)

def build_wp_api_posts_url(url):
    url = (url or '').strip()
    if not url:
        return ''
    parsed = urlparse(url)
    base = f"{parsed.scheme}://{parsed.netloc}" if parsed.scheme and parsed.netloc else url.rstrip('/')
    normalized = url.rstrip('/')

    if '/wp-json/' in normalized:
        if normalized.endswith('/wp/v2/posts'):
            return normalized
        if normalized.endswith('/wp-json/wp/v2'):
            return normalized + '/posts'
        if normalized.endswith('/wp-json'):
            return normalized + '/wp/v2/posts'

    return base.rstrip('/') + '/wp-json/wp/v2/posts'

def format_wp_api_date(date_str):
    if not date_str:
        return "unknown-date"
    try:
        dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        return dt.strftime('%Y-%m-%d')
    except ValueError as e:
        log.error(f"Error parsing API date '{date_str}': {e}")
        return "unknown-date"

def fetch_wp_posts_page(api_posts_url, headers, page=1, per_page=WP_API_MAX_PER_PAGE):
    params = {
        'per_page': per_page,
        'page': page,
        'orderby': 'date',
        'order': 'desc',
        'status': 'publish'
    }
    response = requests.get(api_posts_url, headers=headers, params=params)
    response.raise_for_status()
    posts = response.json()
    total_pages = int(response.headers.get('X-WP-TotalPages', 1))
    return posts, total_pages
def download_blog_posts_wp_api(url, number=None, outdir='posts', word=False, indir=None):
    '''
    Download blog posts using the WordPress REST API instead of scraping HTML pages.
    '''
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
    }

    # Set of already downloaded posts based on indir_list
    existing_txt_files = set()
    existing_docx_files = set()
    indir_list = indir if indir is not None else []
    for idir in (indir_list or []):
        if os.path.isdir(idir):
            for fname in os.listdir(idir):
                if fname.endswith(".txt"):
                    existing_txt_files.add(fname)
                elif fname.endswith(".docx"):
                    existing_docx_files.add(fname)

    if not os.path.exists(outdir):
        os.makedirs(outdir)

    # Track already-seen URLs
    seen_urls_path = os.path.join(outdir, 'seen_urls.txt')
    seen_urls = set()
    if os.path.exists(seen_urls_path):
        with open(seen_urls_path, 'r') as f:
            seen_urls = set(line.strip() for line in f if line.strip())

    api_posts_url = build_wp_api_posts_url(url)
    log.info(f'WordPress API endpoint: {api_posts_url}')

    count = 0
    page = 1
    per_page = min(number, WP_API_MAX_PER_PAGE) if number else WP_API_MAX_PER_PAGE
    consecutive_duplicates = 0

    while True:
        try:
            posts, total_pages = fetch_wp_posts_page(api_posts_url, headers, page=page, per_page=per_page)
        except requests.RequestException:
            log.error(f"Failed to fetch WordPress API page: {api_posts_url} (page {page})")
            return

        if not posts:
            log.warning("No posts returned by WordPress API.")
            break

        for post in posts:
            if number and count >= number:
                return

            post_url = post.get('link') or post.get('guid', {}).get('rendered')
            if post_url and post_url in seen_urls:
                log.info(f"Skipping already-seen URL: {post_url}")
                continue

            post_title_html = (post.get('title', {}) or {}).get('rendered', '')
            post_title = BeautifulSoup(post_title_html, 'html.parser').get_text().strip()
            if not post_title:
                log.warning("No title found in API post, skipping...")
                continue

            publication_date = format_wp_api_date(post.get('date') or post.get('date_gmt'))
            content_html = (post.get('content', {}) or {}).get('rendered', '')
            content_text = extract_text_from_html(content_html, outdir)

            if not content_text:
                log.warning(f"No content found in {post_title}, skipping...")
                continue

            sanitized_title = sanitize_filename(post_title)
            filename_with_date = f"{sanitized_title}-{publication_date}"
            text_filename = filename_with_date + '.txt'
            docx_filename = filename_with_date + '.docx'

            if is_duplicate_filename(text_filename, docx_filename, existing_txt_files, existing_docx_files):
                log.info(f"Skipping duplicate: {text_filename}")
                consecutive_duplicates += 1
                if consecutive_duplicates >= NUM_DUPLICATES:
                    log.info(f"Found {NUM_DUPLICATES} filename-based duplicates in a row. Stopping early.")
                    return
                continue
            else:
                consecutive_duplicates = 0

            save_text_file(post_title, publication_date, content_text, post_url or '', outdir, filename_with_date)
            save_word_file(post_title, publication_date, content_text, outdir, filename_with_date, word, log, post_url or '')

            if post_url:
                with open(seen_urls_path, 'a') as f:
                    f.write(post_url + '\n')

            count += 1

        if page >= total_pages:
            log.info('No more API pages to process.')
            break
        page += 1

def extract_date_from_filename(filename):
    """
    Extracts and returns the date from a given filename in the format '[filename]-YYYY-MM-DD'.
    Returns None if the date format is invalid.
    """
    try:
        parts = filename.split('-')
        if len(parts) >= 3:
            date_str = '-'.join(parts[-3:])  # Extract the last three parts for YYYY-MM-DD
            date_str = date_str.split('.')[0]  # Remove file extension
            return datetime.strptime(date_str, "%Y-%m-%d").date()
    except ValueError:
        log.warning(f"Invalid date format in filename: {filename}")
    return None
            
def concatenate_txt_files(indir_list, outfile, num_files):
    '''
    Concatenate text files in the specified directories, sorted by date (newest first).
    Assumes that a function extract_date_from_filename(filename) is defined externally.
    '''
    files_with_dates = []
    for idir in indir_list:
        if not os.path.isdir(idir):
            continue
        for file in os.listdir(idir):
            if file.endswith('.txt'):
                file_date = extract_date_from_filename(file)
                if file_date:
                    files_with_dates.append((os.path.join(idir, file), file_date))
                else:
                    log.warning(f"Filename does not contain a valid date: {file}")

    # Sort files by date (newest first)
    files_with_dates.sort(key=lambda x: x[1], reverse=True)

    # Process files in sorted order
    with open(f'{outfile}.txt', 'w') as out_file:
        processed_files = 0
        for file_path, _ in files_with_dates:
            with open(file_path, 'r') as in_file:
                out_file.write(in_file.read() + '\n\n')
                log.info(f'Processed file: {os.path.basename(file_path)}')
                processed_files += 1

                # Break if num_files is reached
                if num_files is not None and processed_files >= num_files:
                    break

    log.info(f'Concatenated document saved as: {outfile}.txt')


def concatenate_word_documents(indir_list, outfile, num_files=None):
    '''
    Concatenates Word documents from the specified directories into a single document.
    Properly preserves headings, paragraphs, and images.
    '''
    files_with_dates = []

    # Collect and sort files by date (newest first)
    for idir in indir_list:
        if not os.path.isdir(idir):
            continue
        for file in os.listdir(idir):
            if file.endswith('.docx'):
                file_date = extract_date_from_filename(file)
                if file_date:
                    files_with_dates.append((os.path.join(idir, file), file_date))
                else:
                    log.warning(f"Filename does not contain a valid date: {file}")

    # Sort files by date (newest first)
    files_with_dates.sort(key=lambda x: x[1], reverse=True)

    # Process files in sorted order
    outdoc = Document()
    processed_files = 0

    for file_path, _ in files_with_dates:
        src_doc = Document(file_path)

        # Add title as a heading (use filename without the date)
        title_text = os.path.basename(file_path).rsplit('-', 1)[0].replace('_', ' ')
        outdoc.add_heading(title_text, level=1)

        # Copy paragraphs and formatting from the original document
        i = 0
        while i < len(src_doc.paragraphs):
            paragraph = src_doc.paragraphs[i]
            # Preserve plain text of link instead of hyperlink
            if paragraph.text.strip() == "Original post link":
                if i + 1 < len(src_doc.paragraphs):
                    link_url = src_doc.paragraphs[i + 1].text.strip()
                    log.debug(f"Post link: {link_url}")
                    outdoc.add_paragraph(f"Original post link: {link_url}")
                i += 2
                continue
            new_paragraph = outdoc.add_paragraph(style=paragraph.style)
            for run in paragraph.runs:
                r = new_paragraph.add_run(run.text)
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline
            i += 1

        # Insert a page break between posts
        outdoc.add_page_break()

        log.info(f'Processed Word file: {os.path.basename(file_path)}')
        processed_files += 1

        if num_files is not None and processed_files >= num_files:
            break

    # Save the concatenated document
    outdoc.save(f'{outfile}.docx')
    log.info(f'Concatenated document saved as: {outfile}.docx')

    
# ****************************************************************************************
# arguments
# ****************************************************************************************
def handle_args():
    parser = argparse.ArgumentParser(description='Download blog posts as text documents')
    parser.add_argument(
        '--url', 
        help='URL of the blog page')
    parser.add_argument(
        '--number', 
        type=int, 
        help='Number of blog posts to download')
    parser.add_argument(
        '--outdir', 
        default='posts', 
        help='Directory to save blog posts [default: posts]')    
    parser.add_argument(
        '--indir',
        nargs='+',
        help='One or more input dirs for file concat (space-separated)')
    parser.add_argument(
        '--outfile', 
        default='all_files', 
        help='Concat file name [default: all_files]')    
    parser.add_argument(
        '--word',
        action='store_true',
        help='Download blog posts as formatted Word documents instead of text files.')
    parser.add_argument(
        '--download',
        action='store_true',
        help='Download blog posts.')    
    parser.add_argument(
        '--concat',
        action='store_true',
        help='Concat files.')
    parser.add_argument(
        '-v',
        '--verbose',
        action='store_true',
        help='Enable verbose output to stdout.')
    parser.add_argument(
        '-q',
        '--quiet',
        action='store_true',
        help='Minimal stdout.')
    args = parser.parse_args()

    # Configure stdout logging based on arguments
    ch = logging.StreamHandler(sys.stdout)
    if args.verbose:
        ch.setLevel(logging.DEBUG)
    elif args.quiet:
        ch.setLevel(logging.ERROR)
    else:
        ch.setLevel(logging.INFO)
    ch.setFormatter(formatter)
    log.addHandler(ch)
    
    # check requirements
    if args.download:
        if args.url:
            log.debug(f'Download requirements met')
        else:
            log.error(f'XXX Must supply a URL')
            sys.exit(1)

    log.info('++++++++++++++++++++++++++++++++++++++++++++++')
    log.info(f'+  {os.path.basename(sys.argv[0])}')
    log.info(f'+  Python Version: {sys.version.split()[0]}')
    log.info(f'+  Today is: {date.today()}')
    if args.download:
        log.info(f'+  Target URL: {args.url}')
        log.info(f'+  Number of posts to download: {"All" if args.number is None else args.number}')
        log.info(f'+  Output directory: {args.outdir}')
    if args.concat:
        log.info(f'+  Concatenating files')
        log.info(f'+  Output file: {args.outfile}')
        if not args.download:
            if args.indir:
                log.info(f'+  Concatenating from: {args.indir}')
            else:
                log.error(f'XXX Must include --indir')
                sys.exit(1)
    log.info('++++++++++++++++++++++++++++++++++++++++++++++')        

    return args

# ****************************************************************************************
# Main
# ****************************************************************************************
def main():
    args = handle_args()
    # Build initial list of directories to check for existing posts
    indir_list = args.indir or []
    # Download, using existing indices from all provided indir_list
    if args.download:
        download_blog_posts_wp_api(args.url, args.number, args.outdir, args.word, indir_list)
        # After downloading into outdir, include it for concat
        indir_list = indir_list + [args.outdir]
    # Concatenate from all directories (original and newly downloaded)
    if args.concat:
        concatenate_txt_files(indir_list, args.outfile, args.number)
        if args.word:
            concatenate_word_documents(indir_list, args.outfile, args.number)

if __name__ == '__main__':
    main()
